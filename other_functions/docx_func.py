# %%
'''
docx格式文档处理信息
'''

#%% [markdown]
# ### 整理定点样品

# %%
from typing import Union, List
from copy import deepcopy
import math
from pandas.core.frame import DataFrame

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# %%
def gernerate_record_table_files(
        _info_dict: dict,
        _df: DataFrame,
        _template_file_path: str,
        _output_path: str,
        _tag: Union[str, None] #str | None = None
    )-> None:
    '''
    目的:
        将定点样品编号和相关信息写入记录表模板中
    参数:
        _info_dict:          基本信息
        _df:                 样品编号信息dataframe
        _template_file_path: 记录表模板的路径
        _output_path:        记录表的输出路径
        _tag:                采样的信息标签
    返回:
        无
    '''
    document = Document(_template_file_path)
    table_pages: int = math.ceil((len(_df) - 42) / 24 + 2) # 计算需要多少了记录表的页数
    # 按照需要的表格数量，在docx文档里增减表格
    if table_pages == 1:
        rm_table = document.tables[2]
        t = rm_table._element
        t.getparent().remove(t)
    elif table_pages == 2:
        pass
    else:
        for _ in range(table_pages - 2):
            copy_table = document.tables[2]
            new_table = deepcopy(copy_table)
            _paragraph = document.add_paragraph()
            _paragraph._p.addnext(new_table._element)
    # 根据表格页码，选择Dataframe的范围
    tables = document.tables
    for table_page in range(table_pages):
        if table_page == 0:
            index_first: int = 0
            index_last: int = 17
        else:
            index_first: int = 24 * table_page - 6
            index_last: int = 24 * table_page + 17
        current_df: DataFrame = _df.query(f'index >= {index_first} and index <= {index_last}')
        # 向指定表格填写数据
        current_table = tables[table_page + 1]
        for r_i in range(len(current_df)):
            for c_i in range(3):
                current_cell = current_table.rows[r_i + 2].cells[c_i]
                current_cell.text = str(current_df.iloc[r_i, c_i])
                
                if c_i <=1:
                    current_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                else:
                    current_cell.paragraphs[0].runs[0].font.size = Pt(6.5)
    # 删除多余的换行符
    rm_p = document.paragraphs[2]
    p = rm_p._element
    p.getparent().remove(p)
    # 添加基本信息
    info_table = tables[0]
    code_cell = info_table.rows[0].cells[1]
    comp_cell = info_table.rows[0].cells[3]
    item_cell = info_table.rows[3].cells[1]
    code_cell.text = _info_dict["code"]
    comp_cell.text = _info_dict["comp"]
    item_cell.text = _info_dict["item"]
    # 基本信息单元格样式
    for cell in [code_cell, comp_cell, item_cell]:
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
        p.runs[0].font.size = Pt(9)
    
    output_file_path: str = f'{_output_path}\\{_info_dict["code"]}--{_tag}--{_info_dict["item"]}'
    document.save(f"{output_file_path}.docx")

# %% 将定点样品编号写入记录表模板中

def handle_point_samples(
        _info_dict: dict,
        _point_df: DataFrame,
        _template_file_path: str,
        _output_path_str: str,
        _tag: Union[str, None]
) -> None:
    '''
    目的:
        将定点样品编号和相关信息写入记录表模板中
    参数:
        _info_dict:          基本信息
        _point_df:           定点样品编号信息dataframe
        _template_file_path: 记录表模板的路径
        _output_path:        记录表的输出路径
        _tag:                采样的信息标签
    返回:
        无
    '''

    # 合并相应列，有时可能不需要一部分列，需要调整
    _point_df["工作岗位"] = _point_df["车间/单元"] + "\n" + _point_df["检测地点/岗位"]# + "\n" + _point_df["工种"]
    # _point_df["工作岗位"] = f'{_point_df["车间/单元"]}\n{_point_df["检测地点/岗位"]}'
    _new_df: DataFrame = _point_df[["采样点编号", "工作岗位", "样品编号", "检测项目"]]
    _items_list: List[str] = _new_df['检测项目'].drop_duplicates().tolist()
    for item in _items_list:
        _info_dict["item"] = item
        _current_df: DataFrame = _new_df.query("检测项目 == @item").reset_index(drop=True)
        gernerate_record_table_files(
            _info_dict,
            _current_df,
            _template_file_path,
            _output_path_str,
            _tag
        )
