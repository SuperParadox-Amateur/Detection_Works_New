'''
重构，每个函数或者方法都可以直接出结果。考虑使用functools模块
# [x] 考虑将操作word文件的库改为pywin32
'''
#%%
from io import BytesIO
import math
import os
import re
from copy import deepcopy
from decimal import Decimal, ROUND_HALF_UP
from typing import Any, Dict, List, Tuple, Optional
from nptyping import DataFrame  # , Structure as S
import numpy as np
import pandas as pd
from pandas.api.types import CategoricalDtype
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# from occupational_health_module.other_infos import templates_info

#%%
templates_info: Dict[str, Dict[str, Any]] = {
    '有害物质定点': {
        'template_path': './templates/有害物质定点采样记录.docx',
        # 'template_doc': Document('./templates/有害物质定点采样记录.docx'),
        'direct-reading': False,
        'join_char': '\n',
        'project_num_row': 0,
        'project_num_col': 1,
        'company_name_row': 0,
        'company_name_col': 4,
        'deleterious_substance_row': 3,
        'deleterious_substance_col': 1,
        'first_page_rows': 18,
        'late_page_rows': 24,
        'title_rows': 2,
        'item_rows': 6,
        'available_cols': [0, 1, 2, 4]
    },
    '有害物质个体': {
        'template_path': './templates/有害物质个体采样记录.docx',
        # 'template_doc': Document('./templates/有害物质个体采样记录.docx'),
        'direct-reading': False,
        'join_char': '\n',
        'project_num_row': 0,
        'project_num_col': 1,
        'company_name_row': 0,
        'company_name_col': 4,
        'deleterious_substance_row': 3,
        'deleterious_substance_col': 1,
        'first_page_rows': 5,
        'late_page_rows': 6,
        'title_rows': 2,
        'item_rows': 3,
        'available_cols': [0, 1, 2, 4]
    },
    '噪声定点': {
        'template_path': './templates/噪声定点采样记录.docx',
        # 'template_doc': Document('./templates/噪声定点采样记录.docx'),
        'direct-reading': True,
        'join_char': ' ',
        'project_num_row': 0,
        'project_num_col': 1,
        'company_name_row': 1,
        'company_name_col': 1,
        'deleterious_substance_row': 0,
        'deleterious_substance_col': 0,
        'first_page_rows': 10,
        'late_page_rows': 12,
        'title_rows': 2,
        'item_rows': 1,
        'available_cols': [0, 1, 2]
    },
    '噪声个体': {
        'template_path': './templates/噪声个体采样记录.docx',
        # 'template_doc': Document('./templates/噪声个体采样记录.docx'),
        'direct-reading': True,
        'join_char': ' ',
        'project_num_row': 0,
        'project_num_col': 1,
        'company_name_row': 1,
        'company_name_col': 1,
        'deleterious_substance_row': 0,
        'deleterious_substance_col': 0,
        'first_page_rows': 9,
        'late_page_rows': 11,
        'title_rows': 2,
        'item_rows': 1,
        'available_cols': [0, 1, 2]
    },
    '高温定点': {
        'template_path': './templates/高温定点采样记录.docx',
        # 'template_doc': Document('./templates/高温定点采样记录.docx'),
        'direct-reading': True,
        'join_char': '\n',
        'project_num_row': 0,
        'project_num_col': 1,
        'company_name_row': 1,
        'company_name_col': 1,
        'deleterious_substance_row': 0,
        'deleterious_substance_col': 0,
        'first_page_rows': 1,
        'late_page_rows': 2,
        'title_rows': 3,
        'item_rows': 9,
        'available_cols': [0, 1]
    },
    '一氧化碳定点': {
        'template_path': './templates/一氧化碳定点采样记录.docx',
        # 'template_doc': Document('./templates/一氧化碳定点采样记录.docx'),
        'direct-reading': True,
        'join_char': '\n',
        'project_num_row': 0,
        'project_num_col': 1,
        'company_name_row': 0,
        'company_name_col': 3,
        'deleterious_substance_row': 0,
        'deleterious_substance_col': 0,
        'first_page_rows': 20,
        'late_page_rows': 20,
        'title_rows': 2,
        'item_rows': 4,
        'available_cols': [0, 1]
    },
}

#%%
# [x] 判断系统使用Word或WPS Office
# import win32com.client

# def is_word_installed():
#    try:
#        word_app = win32com.client.Dispatch("Word.Application")
#        word_app.Quit()
#        return True
#    except:
#        return False

# def is_wps_installed():
#    try:
#        wps_app = win32com.client.Dispatch("kwps.Application")
#        wps_app.Quit()
#        return True
#    except:
#        return False

# if is_word_installed():
#    print("Microsoft Word已安装")
# elif is_wps_installed():
#    print("WPS Office已安装")
# else:
#    print("未检测到Word或WPS Office安装")

class OccupationalHealthItemInfo():
    '''职业卫生相应信息生成'''

    def __init__(
            self,
            company_name: str,
            project_number: str,
            point_info_df: DataFrame,
            personnel_info_df: DataFrame,
            types_order: List[str] = ['空白', '定点', '个体']
    ) -> None:
        self.company_name: str = company_name
        self.project_number: str = project_number
        self.templates_info: Dict = templates_info
        self.default_types_order: List[str] = types_order
        # self.default_types_order: List[str] = ['空白', '定点', '个体']
        self.point_info_df: DataFrame = self.initialize_point_df(point_info_df)
        self.personnel_info_df: DataFrame = self.initialize_personnel_df(personnel_info_df)
        self.single_day_engaged_num: Dict[str, int] = {
            '0': 0
        }
        self.output_path: str = os.path.join(
            os.path.expanduser("~/Desktop"),
            f'{self.project_number}记录信息'
        )
        # [x] 数据预先操作方法
        self.factor_reference_df: DataFrame = self.get_occupational_health_factor_reference()
        self.sort_df()
        self.get_detection_days()
        self.schedule_days: int = self.get_schedule_days()  # 采样日程总天数
        (
            self.point_deleterious_substance_df,
            self.personnel_deleterious_substance_df
        ) = self.get_deleterious_substance_df()
        self.output_deleterious_substance_info_dict: Dict = self.get_all_days_dfs()

    # [x] 增加转换相应列为对应数据类型的方法
    def initialize_point_df(self, point_df) -> DataFrame:
        '''转换定点信息df的数据类型'''
        point_dtypes = {
            '采样点编号': int,
            '单元': str,
            '检测地点': str,
            '工种': str,
            '日接触时间': float,
            '检测因素': str,
            '采样数量/天': int,
            '采样日程': str,
        }
        new_point_df = (
            point_df.astype(point_dtypes)
        )
        return new_point_df

    def initialize_personnel_df(self, personnel_df) -> DataFrame:
        '''转换个体信息df的数据类型'''
        personnel_dtypes = {
            '采样点编号': str,
            '单元': str,
            '工种': str,
            '日接触时间': float,
            '检测因素': str,
            '采样数量/天': int,
            '采样日程': str,
        }
        new_personnel_df = (
            personnel_df.astype(personnel_dtypes)
        )
        return new_personnel_df


    # [x] 创建默认的保存路径
    def create_normal_folder(self) -> None:
        '''创建默认的保存路径'''
        if not os.path.exists(self.output_path):
            os.mkdir(self.output_path)
        else:
            pass

    # [x] 默认获得职业卫生所有检测因素的参考信息
    def get_occupational_health_factor_reference(self) -> DataFrame:
        '''
        获得职业卫生所有检测因素的参考信息
        '''
        reference_path: str = os.path.join(
            # self.upper_abs_path,
            'info_files/检测因素参考信息.csv'
        )
        reference_df: DataFrame = pd.read_csv(reference_path)  # type: ignore
        # [x] 增加不同列的空值为不同的数值
        reference_df: DataFrame = reference_df.fillna('/')
        return reference_df

    # [x] 检测信息排序
    def sort_df(self) -> None:
        '''
        对检测信息里的检测因素进行排序
        先对复合检测因素内部排序，再对所有检测因素
        '''
        # 将检测因素，尤其是复合检测因素分开转换为列表，并重新排序
        factor_reference_list: List[str] = (
            self.factor_reference_df['标识检测因素']
            .tolist()
        )
        self.point_info_df['检测因素'] = (
            self
            .point_info_df['检测因素']
            .str.split('|')
            # type: ignore
            .apply(self.custom_sort, args=(factor_reference_list,))
        )
        self.personnel_info_df['检测因素'] = (
            self
            .personnel_info_df['检测因素']
            .str.split('|')
            # type: ignore
            .apply(self.custom_sort, args=(factor_reference_list,))
        )
        # 将检测因素列表的第一个作为标识
        self.point_info_df['标识检测因素'] = (
            self
            .point_info_df['检测因素']
            .apply(lambda lst: lst[0])  # type: ignore
        )
        self.personnel_info_df['标识检测因素'] = (
            self
            .personnel_info_df['检测因素']
            .apply(lambda lst: lst[0])  # type: ignore
        )
        # 将检测因素列表合并为字符串
        self.point_info_df['检测因素'] = (
            self
            .point_info_df['检测因素']
            .apply(lambda x: '|'.join(x))   # type: ignore
        )
        self.personnel_info_df['检测因素'] = (
            self
            .personnel_info_df['检测因素']
            .apply(lambda x: '|'.join(x))   # type: ignore
        )
        # 将定点和个体的检测因素提取出来，创建CategoricalDtype数据，按照拼音排序
        point_factor_list: List[str] = (
            self
            .point_info_df['检测因素']
            .unique().tolist()  # type: ignore
        )
        point_factor_list: List[str] = sorted(
            point_factor_list,
            key=lambda x: x.encode('gbk')
        )
        point_factor_order: CategoricalDtype = CategoricalDtype(
            point_factor_list, ordered=True)
        personnel_factor_list: List[str] = (
            self
            .personnel_info_df['检测因素']
            .unique().tolist()  # type: ignore
        )
        personnel_factor_list: List[str] = sorted(
            personnel_factor_list,
            key=lambda x: x.encode('gbk')
        )
        personnel_factor_order: CategoricalDtype = CategoricalDtype(
            personnel_factor_list,
            ordered=True
        )
        # 将检测因素按照拼音排序
        self.point_info_df['检测因素'] = (
            self
            .point_info_df['检测因素']
            .astype(point_factor_order)  # type: ignore
        )
        self.point_info_df = (
            self
            .point_info_df
            # type: ignore
            .sort_values(by=['检测因素', '采样点编号'], ascending=True, ignore_index=True)
        )
        self.personnel_info_df['检测因素'] = (
            self
            .personnel_info_df['检测因素']
            .astype(personnel_factor_order)  # type: ignore
        )
        self.personnel_info_df = (
            self
            .personnel_info_df
            # type: ignore
            .sort_values(by=['检测因素', '采样点编号'], ascending=True, ignore_index=True)
        )

    # [x] 获得采样日程下每一天的检测信息，确定将采样日程改为“1|2|3”或者指定的“2”的方式
    def get_detection_days(self) -> None:
        '''
        获得采样日程下每一天的检测信息
        '''
        self.point_info_df['采样日程'] = (
            self
            .point_info_df['采样日程']
            .str.split('|')  # type: ignore
        )
        self.point_info_df = (
            self
            .point_info_df
            .explode('采样日程', ignore_index=True)
        )
        self.point_info_df['采样日程'] = (
            self
            .point_info_df['采样日程']
            .astype(int)  # type: ignore
        )
        self.personnel_info_df['采样日程'] = (
            self
            .personnel_info_df['采样日程']
            .str.split('|')  # type: ignore
        )
        self.personnel_info_df = (
            self
            .personnel_info_df
            .explode('采样日程', ignore_index=True)
        )
        self.personnel_info_df['采样日程'] = (
            self
            .personnel_info_df['采样日程']
            .astype(int)  # type: ignore
        )

    # [x] 获得所有所有空气有害物质的检测因素，包含定点和个体
    def get_deleterious_substance_df(self) -> Tuple[DataFrame, DataFrame]:
        '''
        获得所有空气有害物质的检测因素，包含定点和个体
        '''
        point_deleterious_substance_df: DataFrame = (
            pd.merge(  # type: ignore
                self.point_info_df,
                self.factor_reference_df[
                    ['标识检测因素', '是否仪器直读', '是否需要空白', '复合因素代码', '收集方式', '定点采样时间']
                ],
                on='标识检测因素',
                how='left'
            )
            .fillna({
                '是否需要空白': False,
                '复合因素代码': 0,
                '是否仪器直读': False,
                '收集方式': '粉尘',
                '定点采样时间': 15,
                })
            .query('是否仪器直读 == False')
        )
        personnel_deleterious_substance_df: DataFrame = (
            pd.merge(  # type: ignore
                self.personnel_info_df,
                self.factor_reference_df[
                    ['标识检测因素', '是否仪器直读', '是否需要空白', '复合因素代码']
                ],
                on='标识检测因素',
                how='left'
            )
            .fillna({
                '是否需要空白': False,
                '复合因素代码': 0,
                '是否仪器直读': False,
            })
            .query('是否仪器直读 == False')
        )
        return point_deleterious_substance_df, personnel_deleterious_substance_df

    # [x] 获得一天的空气有害物质检测因素，包含定点和个体
    def get_single_day_deleterious_substance_df(
        self, schedule_day: int = 1
    ) -> Tuple[DataFrame, DataFrame]:
        '''
        获得一天的空气有害物质检测因素，包含定点和个体
        '''
        single_day_point_deleterious_substance_df: DataFrame = (
            self
            .point_deleterious_substance_df
            [self.point_deleterious_substance_df['采样日程'] == schedule_day]
        )
        single_day_personnel_deleterious_substance_df: DataFrame = (
            self
            .personnel_deleterious_substance_df
            [self.personnel_deleterious_substance_df['采样日程'] == schedule_day]
        )
        single_day_deleterious_substance_df_tuple: Tuple[DataFrame, DataFrame] = (
            single_day_point_deleterious_substance_df,
            single_day_personnel_deleterious_substance_df
        )
        return single_day_deleterious_substance_df_tuple
    # [x] 获得一天的空白样品编号

    def get_single_day_blank_df(self, engaged_num: int = 0, schedule_day: int = 1) -> DataFrame:
        '''
        获得一天的空白样品编号
        '''
        # 应对空白数量为0的情况
        # 复制定点和个体检测信息的dataframe，避免提示错误
        point_df, personnel_df = (
            self
            .get_single_day_deleterious_substance_df(schedule_day)
        )
        single_day_point_df: DataFrame = point_df.copy()
        single_day_personnel_df: DataFrame = personnel_df.copy()
        # 从定点和个体的dataframe提取检测因素，去重以及合并
        single_day_point_df['检测因素'] = (
            single_day_point_df['检测因素']
            .str.split('|')  # type: ignore
        )
        ex_single_day_point_df: DataFrame = (
            single_day_point_df
            .explode('检测因素')
        )
        single_day_personnel_df['检测因素'] = (
            single_day_personnel_df['检测因素']
            .str.split('|')  # type: ignore
        )
        ex_single_day_personnel_df: DataFrame = (
            single_day_personnel_df
            .explode('检测因素')
        )

        # 筛选出需要空白的检测因素
        test_df: DataFrame = (
            pd.concat(
                [
                    ex_single_day_point_df[['检测因素', '是否需要空白', '复合因素代码']],
                    ex_single_day_personnel_df[['检测因素', '是否需要空白', '复合因素代码']]
                ]
            )
            .query('是否需要空白 == True')
            .drop_duplicates('检测因素')
            .reset_index(drop=True)
        )
        # 分别处理非复合因素和复合因素，复合因素要合并。
        # 判断定点和个体的检测因素是否为空
        if test_df.empty:
            single_day_blank_df = pd.DataFrame(columns=['标识检测因素', '空白编号'])
        else:
            raw_group1: DataFrame = test_df.loc[test_df['复合因素代码'] == 0]
            raw_group2: DataFrame = test_df.loc[test_df['复合因素代码'] != 0]
            if raw_group1.empty:
                group1 = pd.DataFrame(columns=['检测因素', '是否需要空白'])
            else:
                group1: DataFrame = raw_group1.loc[:, ['检测因素', '是否需要空白']]
            if raw_group2.empty:
                group2 = pd.DataFrame(columns=['检测因素', '是否需要空白'])
            else:
                group2: DataFrame = (
                    pd.DataFrame(
                        raw_group2.groupby(['复合因素代码'],
                        group_keys=False
                    )
                    ['检测因素']
                    .apply('|'.join))
                    .reset_index(drop=True)
                )
                group2.loc[:, '是否需要空白'] = True

            # 最后合并，排序
            concat_group: DataFrame = pd.concat(  # type: ignore
                [group1, group2],
                ignore_index=True,
                axis=0,
                sort=False
            )
            blank_factor_list: List[str] = sorted(
                concat_group['检测因素'].tolist(),
                key=lambda x: x.encode('gbk')
            )  # type: ignore
            blank_factor_order: CategoricalDtype = CategoricalDtype(
                categories=blank_factor_list,
                ordered=True
            )
            concat_group['检测因素'] = (
                concat_group['检测因素']
                .astype(blank_factor_order)  # type: ignore
            )
            # 筛选出需要空白编号的检测因素，并赋值
            # single_day_blank_df: DataFrame = (
            #     concat_group
            #     # 必须用`==`才可用，按照提示用`is`会失败
            #     .loc[concat_group['是否需要空白'] == True]
            #     .sort_values('检测因素', ignore_index=True)
            # )  # type: ignore
            single_day_blank_df = (
                concat_group
                .assign(
                    是否需要空白=True
                )
                .sort_values('检测因素', ignore_index=True)
            )
            # single_day_blank_df: DataFrame = (
            #     concat_group
            #     # 必须用`==`才可用，按照提示用`is`会失败
            #     .loc[concat_group['是否需要空白'] == True]
            #     .sort_values('检测因素', ignore_index=True)
            # )  # type: ignore
            single_day_blank_df['检测因素'] = (
                single_day_blank_df['检测因素']
                .astype(str)
                # type: ignore
                .map(lambda x: [x] + x.split('|') if x.count('|') > 0 else x)
            )
            single_day_blank_df['空白编号'] = (
                np.arange(1, single_day_blank_df.shape[0] + 1)
                + engaged_num  # type: ignore
            )
            single_day_blank_df.drop(
                columns=['是否需要空白'],
                inplace=True
            )  # type: ignore
            single_day_blank_df = (
                single_day_blank_df
                .explode('检测因素')
                .rename(columns={'检测因素': '标识检测因素'})
            )
        return single_day_blank_df
    # [x] 处理单日的定点检测信息，为其加上样品编号范围和空白样品编号

    def get_single_day_point_df(self, engaged_num: int = 0, schedule_day: int = 1) -> DataFrame:
        '''
        处理单日的定点检测信息，为其加上样品编号范围和空白样品编号
        '''
        # 注：为定点添加空白编号的功能不要放到这里实现
        point_df: DataFrame = (
            self
            .get_single_day_deleterious_substance_df(schedule_day)[0]
            .copy()
        )
        # [x] 如果数量为0
        point_df['终止编号'] = (
            point_df['采样数量/天'].cumsum()
            + engaged_num  # type: ignore
        )
        point_df['起始编号'] = point_df['终止编号'] - point_df['采样数量/天'] + 1
        # [x] 是否合并代表时长
        point_df['是否合并代表时长'] = (
            point_df # type: ignore
            .apply(
                lambda df: True if df['日接触时间'] / df['采样数量/天'] < 0.25 else False,
                axis=1
            ) # type: ignore
        )

        return point_df
    # [x] 处理单日的个体检测信息，为其加上样品编号范围和空白样品编号

    def get_single_day_personnel_df(self, engaged_num: int = 0, schedule_day: int = 1) -> DataFrame:
        '''
        处理单日的个体检测信息，为其加上样品编号范围和空白样品编号
        '''
        personnel_df: DataFrame = (
            self
            .get_single_day_deleterious_substance_df(schedule_day)
            [1].copy()
        )
        # [x] 如果数量为0
        personnel_df['个体编号'] = (
            personnel_df['采样数量/天'].cumsum()
            + engaged_num  # type: ignore
        )
        return personnel_df

    # [x] 获得单日的所有编号排列好的样品信息

    def get_single_day_dfs(self, schedule_day: int = 1) -> Dict[str, DataFrame]:
        '''为单日的监测信息的样品编号'''
        engaged_num_copy: int = (
            self
            .single_day_engaged_num[f'{schedule_day - 1}']
        )
        for type_order in self.default_types_order:
            if type_order == '空白':
                current_blank_df: DataFrame = (
                    self.get_single_day_blank_df(
                        engaged_num_copy,
                        schedule_day
                    )
                )
                engaged_num_copy: int = (
                    self.refresh_engaged_num(
                        current_blank_df,
                        type_order,
                        engaged_num_copy
                    )
                )
            elif type_order == '定点':
                current_point_df: DataFrame = (
                    self.get_single_day_point_df(
                        engaged_num_copy,
                        schedule_day
                    )
                )
                engaged_num_copy: int = (
                    self.refresh_engaged_num(
                        current_point_df,
                        type_order,
                        engaged_num_copy
                    )
                )
            elif type_order == '个体':
                current_personnel_df: DataFrame = (
                    self.get_single_day_personnel_df(
                        engaged_num_copy,
                        schedule_day
                    )
                )
                engaged_num_copy: int = (
                    self.refresh_engaged_num(
                        current_personnel_df,
                        type_order,
                        engaged_num_copy
                    )
                )
        # [x] 更新单日的已占用编号
        self.single_day_engaged_num[f'{schedule_day}'] = engaged_num_copy
        # [x] 为个体添加空白编号
        if current_blank_df.empty == False:  # type: ignore
            current_point_df: DataFrame = (
                pd.merge(
                    current_point_df,  # type: ignore
                    current_blank_df,  # type: ignore
                    how='left',
                    on='标识检测因素'
                )
                .fillna({'空白编号': 0})
            )
            # current_point_df = (
            #     pd.merge(
            #         current_blank_df,
            #         on='标识检测因素',
            #         how='left'
            #     )
            #     .fillna(0)
            # )
        else:
            # current_point_df.loc[:, '空白编号'] = 0  # type: ignore
            current_point_df = current_point_df.assign(空白编号 = 0)  # type: ignore
        # 相应的列转为整数
        int_list: List[str] = ['终止编号', '起始编号', '空白编号']
        current_point_df[int_list] = (  # type: ignore
            current_point_df[int_list]  # type: ignore
            .astype(int)
        )
        # 获得编号列表
        # [x] 爆炸的定点编号
        current_point_df_copy: DataFrame = current_point_df.copy()  # type: ignore
        # 样品编号加入空白编号
        current_point_df_copy['样品编号'] = (  # type: ignore
            current_point_df_copy  # type: ignore
            .apply(
                lambda df: self.get_exploded_point_df(
                    df['空白编号'],
                    df['起始编号'],
                    df['终止编号']
                ),
                axis=1
            )
        )
        # [x] 添加代表时长列（转移到这里）
        # 获得代表时长列表
        current_point_df_copy['代表时长'] = (
            current_point_df_copy
            .apply(
                lambda df: self.get_exploded_contact_duration(
                    df['日接触时间'], df['采样数量/天'], 4
                ),
                axis=1
            )
        )
        # 爆炸定点编号
        ex_current_point_df: DataFrame = (
            current_point_df_copy
            .explode(['样品编号', '代表时长'])  # type: ignore
        )
        # [x] 增加获得流转单信息功能
        counted_df: DataFrame = (
            self
            .get_single_day_dfs_stat(current_point_df, current_personnel_df) # type: ignore
        )
        # 字典保存
        df_dict: Dict[str, DataFrame] = {
            '空白': current_blank_df,  # type: ignore
            '定点': current_point_df,  # type: ignore
            '爆炸定点': ex_current_point_df,
            '个体': current_personnel_df,  # type: ignore
            '样品统计': counted_df,
        }
        return df_dict

    # [x] 获得每日的样品编号信息，并存储到相应字典里
    def get_all_days_dfs(self):
        '''获得每日的样品编号信息，并存储到相应字典里'''
        output_deleterious_substance_dict = {}
        for schedule_day in range(1, self.schedule_days + 1):
            current_df_dict: Dict[str, DataFrame] = (
                self
                .get_single_day_dfs(schedule_day)
            )
            output_deleterious_substance_dict[f'{schedule_day}'] = current_df_dict
        return output_deleterious_substance_dict

    # [x] 获得爆炸后的定点样品编号

    def get_exploded_point_df(
        self,
        blank_num: int,
        start_num: int,
        end_num: int
    ) -> List[Optional[str]]:
        '''将定点df爆炸成多行的定点df'''
        # 空白编号
        if blank_num != 0:
            blank_list: List[Optional[str]] = [
                f'{blank_num:0>4d}-1',
                f'{blank_num:0>4d}-2',
            ]
        else:
            blank_list: List[Optional[str]] = [None, None]
        # 定点编号
        point_list: List[int] = list(range(start_num, end_num + 1))
        point_str_list: List[Optional[str]] = [
            f'{i:0>4d}' for i in point_list
        ]
        point_str_list_extra: List[Optional[str]] = [None] * (4 - len(point_str_list))
        point_str_list.extend(point_str_list_extra)
        # 空白加定点
        all_list: List[Optional[str]] = blank_list + point_str_list

        return all_list

    # [x] 获得分开的接触时间，使用十进制来计算

    def get_exploded_contact_duration(
        self,
        duration: float,
        size: int,
        full_size: int
    ) -> List[Optional[float]]:
        '''获得分开的接触时间，使用十进制来计算'''
        time_dec: Decimal = Decimal(str(duration))
        size_dec: Decimal = Decimal(str(size))
        time_list_dec: List[Decimal] = []
        if time_dec < Decimal('0.25') * size_dec:
            time_list_dec.append(time_dec)
        elif time_dec < Decimal('0.3') * size_dec:
            front_time_list_dec: List[Decimal] = [
                Decimal('0.25')] * (int(size) - 1)
            last_time_dec: Decimal = time_dec - sum(front_time_list_dec)
            time_list_dec.extend(front_time_list_dec)
            time_list_dec.append(last_time_dec)
        else:
            time_prec: int = int(time_dec.as_tuple().exponent)
            if time_prec == 2:
                prec_str: str = '0.00'
            else:
                prec_str: str = '0.0'
            judge_result: Decimal = time_dec / size_dec
            for _ in range(int(size) - 1):
                result: Decimal = judge_result.quantize(
                    Decimal(prec_str), ROUND_HALF_UP)
                time_list_dec.append(result)
            last_result: Decimal = time_dec - sum(time_list_dec)
            time_list_dec.append(last_result)

        time_list: List[float] = sorted(
            list(map(float, time_list_dec)), reverse=False)
        # str_time_list: List[str] = list(map(str, time_list))
        blank_cell_list: List[None] = [None, None]
        complement_cell_list: List[None] = [None] * (full_size - len(time_list))
        all_time_list: List[Optional[float]] = (
            blank_cell_list
            + time_list
            + complement_cell_list
        )

        return all_time_list
    # [x] 整理定点和个体的样品统计信息

    def get_single_day_dfs_stat(
        self,
        current_point_df: DataFrame,
        current_personnel_df: DataFrame
    ) -> DataFrame:
        '''整理定点和个体的样品信息'''
        pivoted_point_df: DataFrame = (
            pd.pivot_table(
                current_point_df,
                index=['检测因素'],
                aggfunc={'空白编号': max, '起始编号': min, '终止编号': max}
            )
        )
        # 增加个体样品数量为0时的处理方法
        # [x] 增加空白样品数量为0时的处理方法
        if current_personnel_df.empty == False:
            pivoted_personnel_df: DataFrame = (
                pd.pivot_table(
                    current_personnel_df,
                    index=['检测因素'],
                    values='个体编号',
                    aggfunc=[min, max]
                )
                .stack()
                .reset_index()
                .set_index('检测因素')
                .drop('level_1', axis=1)
                .rename(columns={'min': '个体起始编号', 'max': '个体终止编号'})
            )
        else:
            pivoted_personnel_df = pd.DataFrame(columns=['个体起始编号', '个体终止编号'])
            pivoted_personnel_df.index.name = '检测因素'

        # 合并空白、定点和个体的信息
        counted_df: DataFrame = (
            pd.concat([pivoted_point_df, pivoted_personnel_df], axis=1)
            .fillna(0)
            .applymap(int)
        )
        # 统计空白、定点和个体的数量
        counted_df['空白数量'] = (
            counted_df['空白编号']
            .apply(lambda x: 2 if x != 0 else 0)
        )
        counted_df['定点数量'] = (
            counted_df
            .apply(
                lambda x: x['终止编号'] - x['起始编号'] + 1 if x['终止编号'] != 0 else 0,
                axis=1
            )
        )
        counted_df['个体数量'] = (
            counted_df
            .apply(
                lambda x: x['个体终止编号'] - x['个体起始编号'] +
                1 if x['个体终止编号'] != 0 else 0,
                axis=1
            )
        )
        counted_df['总计'] = (
            counted_df['空白数量']
            + counted_df['定点数量']
            + counted_df['个体数量']
        )
        # 统计空白、定点和个体的编号范围
        counted_df['空白编号范围'] = (
            counted_df
            .apply(
                self.get_blank_count_range,
                axis=1
            )
        )
        counted_df['定点编号范围'] = (
            counted_df
            .apply(
                self.get_point_count_range,
                axis=1
            )
        )
        counted_df['个体编号范围'] = (
            counted_df
            .apply(
                self.get_personnel_count_range,
                axis=1
            )
        )
        counted_df['编号范围'] = (
            self.project_number
            + counted_df
            .apply(self.get_range_str, axis=1)
        )
        counted_df['检测因素c'] = counted_df.index
        counted_df['保存时间'] = (
            counted_df['检测因素c']
            .apply(self.get_counted_df_save_info)
        )
        counted_df = counted_df.reset_index(drop=False)
        return counted_df

    # [x] 将每日空白信息，定点编号，爆炸定点编号，个体编号和样品统计信息写入excel文件里
    def write_output_deleterious_substance_info(self) -> None:
        '''将每日空白信息，定点编号，爆炸定点编号，个体编号和样品统计信息写入excel文件里'''
        # 缓存到bytes中
        file_io: BytesIO = BytesIO()
        with pd.ExcelWriter(file_io) as excel_writer:  # pylint: disable=abstract-class-instantiated
            for schedule_day in range(1, self.schedule_days + 1):
                current_output_info_dict: Dict[str, DataFrame] = (
                    self
                    .output_deleterious_substance_info_dict
                    [f'{schedule_day}']
                )
                for name, current_df in current_output_info_dict.items():
                    # 工作表名称
                    sheet_name: str = f'D{schedule_day}{name}'
                    # 保留的列
                    trim_cols: List[str] = self.trim_output_df(name)
                    # 是否保留索引
                    # if name == '样品统计':
                    #     is_index: bool = True
                    # else:
                    #     is_index: bool = False
                    trim_df: DataFrame = current_df[trim_cols]
                    trim_df.to_excel(
                        excel_writer,
                        sheet_name=f'{sheet_name}',
                        index=False
                    )
        file_name: str = f'{self.project_number}-{self.company_name}样品信息.xlsx'
        output_file_path: str = os.path.join(f'{self.output_path}', file_name)
        with open(output_file_path, 'wb') as output_file:
            output_file.write(file_io.getvalue())

    # [x] 整理输出的df

    def trim_output_df(self, name) -> List[str]:
        '''整理输出的df'''
        blank_cols: List[str] = [
            '标识检测因素',
            '空白编号'
        ]
        point_cols: List[str] = [
            '采样点编号',
            '单元',
            '检测地点',
            '工种',
            '日接触时间',
            '检测因素',
            '采样数量/天',
            '采样日程',
            '空白编号',
            '起始编号',
            '终止编号',
        ]
        ex_point_cols: List[str] = [
            '采样点编号',
            '单元',
            '检测地点',
            '工种',
            '日接触时间',
            '检测因素',
            '采样数量/天',
            '采样日程',
            '样品编号',
            '代表时长',
        ]
        personnel_cols: List[str] = [
            '采样点编号',
            '单元',
            '工种',
            '日接触时间',
            '检测因素',
            '采样数量/天',
            '采样日程',
            '个体编号',
        ]
        counted_cols: List[str] = [
            '检测因素',
            '空白编号范围',
            '定点编号范围',
            '个体编号范围',
            '编号范围',
            '总计',
            '保存时间',
        ]
        trim_cols_dict: Dict[str, List[str]] = {
            '空白': blank_cols,
            '定点': point_cols,
            '爆炸定点': ex_point_cols,
            '个体': personnel_cols,
            '样品统计': counted_cols,
        }
        trim_cols: List[str] = trim_cols_dict[name]
        return trim_cols

    # [x] 将信息写入记录表模板里
    def write_to_templates(self):
        '''将信息写入记录表模板里'''
        # 创建文件夹
        self.create_normal_folder()
        self.write_output_deleterious_substance_info()
        # 循环读取天数
        for schedule_day in range(1, self.schedule_days + 1):
            # [x] 定点有害物质
            doc1 = Document(
                self.templates_info['有害物质定点']
                ['template_path']
            )
            self.write_point_deleterious_substance(doc1, schedule_day)
            # [x] 个体有害物质
            doc2 = Document(
                self.templates_info['有害物质个体']
                ['template_path']
            )
            self.write_personnel_deleterious_substance(doc2, schedule_day)
            # [x] 流转单
            traveler_doc = Document(
                self.templates_info['流转单']
                ['template_path']
            )
            self.write_traveler_docx(traveler_doc, schedule_day)
        # [x] 定点仪器直读物质
        other_factors: List[str] = ["一氧化碳", "噪声", "高温"]
        # 不同检测因素调用不同方法处理
        for factor in other_factors:
            # 判断是否存在再调用相应方法处理
            factor_exists: bool = (
                self
                .point_info_df['检测因素']
                .isin([f'{factor}'])
                .any(bool_only=True)
            )
            if factor_exists:
                self.write_direct_reading_factors_docx(factor)
        # [x] 个体噪声
        doc3 = Document(
            self.templates_info['噪声个体']
            ['template_path']
        )
        self.write_personnel_noise(doc3)

    def write_point_deleterious_substance(self, doc: Any, day_i: int) -> None:
        '''将定点有害物质信息写入模板'''
        # merger = PdfWriter()
        # for day_i, schedule in enumerate(self.schedule_list):
        today_df = (
                self
                .output_deleterious_substance_info_dict
                [f'{day_i}']['定点']
                # .query(f'{self.schedule_col} == @schedule')
                .sort_values(by=['采样点编号'])
                .reset_index(drop=True)
            )
        factors: List[str] = today_df['检测因素'].drop_duplicates().tolist()
        sorted_factors: List[str] = sorted(factors, key=lambda x: x.encode('gbk'))
        # 获得当前检测因素的dataframe
        for factor in sorted_factors:
            # 导入定点模板
            doc_copy = deepcopy(doc)
            # 获得当前检测因素的dataframe
            current_factor_df = (
                today_df[today_df['检测因素'] == factor]
                .sort_values(by='采样点编号')
                .reset_index(drop=True)
            )
            # 计算需要的记录表页数
            table_pages: int = (
                math
                .ceil(
                    (len(current_factor_df) - 6)
                    / 4 + 2
                )
            )
            # 按照页数来增减表格数量
            if table_pages == 1:
                rm_table = doc_copy.tables[2]
                t = rm_table._element
                t.getparent().remove(t)
                rm_page_break = doc_copy.paragraphs[-2]
                pg = rm_page_break._element
                pg.getparent().remove(pg)
                rm_page_break2 = doc_copy.paragraphs[-2]
                pg2 = rm_page_break2._element
                pg2.getparent().remove(pg2)
            elif table_pages == 2:
                pass
            else:
                for _ in range(table_pages - 2):
                    cp_table = doc_copy.tables[2]
                    new_table = deepcopy(cp_table)
                    new_paragraph = doc_copy.add_page_break()
                    new_paragraph._p.addnext(new_table._element)
                    doc_copy.add_paragraph()
            # 确定不同的表格要填入的信息范围
            tables = doc_copy.tables
            for table_page in range(table_pages):
                if table_page == 0:
                    index_first: int = 0
                    index_last: int = 2
                else:
                    index_first: int = 4 * table_page - 1
                    index_last: int = 4 * table_page + 2
                current_df = (
                    current_factor_df
                    .query(f'index >= {index_first} and index <= {index_last}')
                    .reset_index(drop=True)
                )
                # 向指定表格填写数据
                current_table = tables[table_page + 1]
                for r_i in range(current_df.shape[0]):
                    # 样品编号列表
                    point_list: List[int] = list(
                        range(
                            current_df.loc[r_i, '起始编号'],
                            current_df.loc[r_i, '终止编号'] + 1
                        )
                    )
                    # point_str_list: List[Optional[str]] = [
                    #     f'{i:0>4d}' for i in point_list
                    # ]
                    # 代表时长列表
                    duration_list: List[Optional[float]] = (
                        self.get_exploded_contact_duration(
                            current_df.loc[r_i, '日接触时间'],
                            current_df.loc[r_i, '采样数量/天'],
                            4
                        )
                    )
                    row_info = {
                        '采样点编号': current_df.loc[r_i, '采样点编号'],
                        '采样岗位': f"{current_df.loc[r_i, '单元']}\n{current_df.loc[r_i, '检测地点']}",
                        '空白编号1': f'{current_df.loc[r_i, "空白编号"]:>04d}-1',
                        '空白编号2': f'{current_df.loc[r_i, "空白编号"]:>04d}-2',
                        '样品编号': point_list,
                        '代表时长': duration_list,
                        '是否合并代表时长': current_df.loc[r_i, '是否合并代表时长'],
                    }
                    # 采样点编号单元格
                    cell1 = current_table.cell(r_i * 6 + 2, 0)
                    cell1.text = str(row_info['采样点编号'])
                    cell1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    cell1.paragraphs[0].runs[0].font.size = Pt(8)
                    # 采样岗位单元格
                    cell2 = current_table.cell(r_i * 6 + 2, 1)
                    cell2.text = row_info['采样岗位']
                    cell2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    cell2.paragraphs[0].runs[0].font.size = Pt(7.5)
                    #[x] 样品编号加上项目编号前缀
                    # 空白编号单元格，只写入第一行
                    if table_page == 0 and r_i == 0:
                        cell3_1 = current_table.cell(r_i * 6 + 2, 2)
                        cell3_1.text = f"{self.project_number}{row_info['空白编号1']}"
                        cell3_2 = current_table.cell(r_i * 6 + 3, 2)
                        cell3_2.text = f"{self.project_number}{row_info['空白编号2']}"
                        cell3_1.paragraphs[0].runs[0].font.size = Pt(8)
                        cell3_2.paragraphs[0].runs[0].font.size = Pt(8)
                    else:
                        pass
                    # 样品编号单元格
                    for n_i, num in enumerate(row_info['样品编号']):
                        cell4 = current_table.cell(r_i * 6 + n_i + 4, 2)
                        cell4.text = f"{self.project_number}{num:0>4d}"
                        cell4.paragraphs[0].runs[0].font.size = Pt(8)
                    # 代表时长
                    for n_i, duration in enumerate(row_info['代表时长']):
                        cell5 = current_table.cell(r_i * 6 + n_i + 2, 9)
                        if duration != None:
                            cell5.text = str(duration)
                            cell5.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                            cell5.paragraphs[0].runs[0].font.size = Pt(9)
                    # 是否合并代表时长
                    if row_info['是否合并代表时长'] == True:
                        merge_len: int = len(row_info['样品编号'])
                        merge_cell1 = current_table.cell(r_i * 6 + 4, 9)
                        merge_cell2 = current_table.cell(r_i * 6 + merge_len + 3, 9)
                        merge_cell1.merge(merge_cell2)
                    #[x] 样式调整
            #[x] 写入项目基本信息
            info_table = tables[0]
            # 项目编号
            code_cell = info_table.cell(0, 1)
            code_cell.text = self.project_number
            # 单位
            comp_cell = info_table.cell(0, 4)
            comp_cell.text = self.company_name
            # 检测因素
            item_cell = info_table.cell(3, 1)
            item_cell.text = str(factor)
            # 采样日期
            date_cell = info_table.cell(3, 6)
            # if self.schedule_col == '采样/送样日期':
            #     date_cell.text = schedule.strftime("%Y年%m月%d日")
            for cell in [code_cell, comp_cell, item_cell, date_cell]:
                p = cell.paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # type: ignore
                if len(cell.text) >= 14:
                    p.runs[0].font.size = Pt(8)
                else:
                    pass
                    # p.runs[0].font.size = Pt(9)
            # 页脚信息
            core_properties = doc_copy.core_properties
            core_properties.keywords = factor
            # if self.schedule_col == "采样/送样日期":
            #     core_properties.comments  = schedule.strftime(r"%Y/%m/%d")
            # 保存到桌面文件夹里
            file_name = f'D{day_i}-定点-{factor}'
            safe_file_name: str = self.convert_safe_filename(file_name)
            file_output_path: str = os.path.join(
                self.output_path,
                safe_file_name
            )
            doc_copy.save(f'{file_output_path}.docx')
            # saved_file_path: str = f'{file_output_path}.docx'
            # with open(saved_file_path, 'rb') as f:
            #     bytes = f.read()
            # file_bytes = BytesIO(bytes)
            # empty_pdf_file = 'empty.pdf'
            # with open(empty_pdf_file, 'wb') as pdf_file:
            #     pass

    def write_personnel_deleterious_substance(self, doc: Any, day_i: int):
        '''将个体有害物质信息写入模板'''
        # for day_i, schedule in enumerate(self.schedule_list):
        today_df = (
                self
                .output_deleterious_substance_info_dict
                [f'{day_i}']['个体']
                # .query(f'{self.schedule_col} == @schedule')
                .sort_values(by=['采样点编号'])
                .reset_index(drop=True)
            )
        factors: List[str] = today_df['检测因素'].drop_duplicates().tolist()
        # 获得当前检测因素的dataframe
        for factor in factors:
            # 导入定点模板
            doc_copy = deepcopy(doc)
            # 获得当前检测因素的dataframe
            current_factor_df = (
                today_df[today_df['检测因素'] == factor]
                .sort_values(by='采样点编号')
                .reset_index(drop=True)
            )
            # 计算需要的记录表页数
            table_pages: int = (
                math
                .ceil(
                    (len(current_factor_df) - 10)
                    / 6 + 2
                )
            )
            # 按照页数来增减表格数量
            if table_pages == 1:
                rm_table = doc_copy.tables[2]
                t = rm_table._element
                t.getparent().remove(t)
                rm_page_break = doc_copy.paragraphs[-2]
                pg = rm_page_break._element
                pg.getparent().remove(pg)
                rm_page_break2 = doc_copy.paragraphs[-2]
                pg2 = rm_page_break2._element
                pg2.getparent().remove(pg2)
            elif table_pages == 2:
                pass
            else:
                for _ in range(table_pages - 2):
                    cp_table = doc_copy.tables[2]
                    new_table = deepcopy(cp_table)
                    new_paragraph = doc_copy.add_page_break()
                    new_paragraph._p.addnext(new_table._element)
                    doc_copy.add_paragraph()
            # 确定不同的表格要填入的信息范围
            tables = doc_copy.tables
            for table_page in range(table_pages):
                if table_page == 0:
                    index_first: int = 0
                    index_last: int = 4
                else:
                    index_first: int = 6 * table_page - 1
                    index_last: int = 6 * table_page + 4
                current_df = (
                    current_factor_df
                    .query(f'index >= {index_first} and index <= {index_last}')
                    .reset_index(drop=True)
                )
                # 向指定表格填写数据
                current_table = tables[table_page + 1]
                for r_i in range(current_df.shape[0]):
                    row_info = {
                        '采样点编号': current_df.loc[r_i, '采样点编号'],
                        '采样岗位': f"{current_df.loc[r_i, '单元']}\n{current_df.loc[r_i, '工种']}",
                        '样品编号': current_df.loc[r_i, '个体编号'],
                        '代表时长': current_df.loc[r_i, '日接触时间'],
                    }
                    # 采样点编号单元格
                    cell1 = current_table.cell(r_i * 3 + 2, 0)
                    cell1.text = str(row_info['采样点编号'])
                    cell1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    cell1.paragraphs[0].runs[0].font.size = Pt(8)
                    # 采样岗位单元格
                    cell2 = current_table.cell(r_i * 3 + 2, 1)
                    cell2.text = row_info['采样岗位']
                    cell2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    cell2.paragraphs[0].runs[0].font.size = Pt(7.5)
                    #[x] 样品编号加上项目编号前缀
                    # 样品编号单元格
                    cell4 = current_table.cell(r_i * 3 + 2, 2)
                    cell4.text = f"{self.project_number}{row_info['样品编号']:0>4d}"
                    cell4.paragraphs[0].runs[0].font.size = Pt(8)
                    # 代表时长
                    cell5 = current_table.cell(r_i * 3 + 2, 4)
                    cell5.text = str(row_info['代表时长'])
                    cell5.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    cell5.paragraphs[0].runs[0].font.size = Pt(9)
                    #[x] 样式调整
            #[x] 写入项目基本信息
            info_table = tables[0]
            # 项目编号
            code_cell = info_table.cell(0, 1)
            code_cell.text = self.project_number
            # 单位
            comp_cell = info_table.cell(0, 4)
            comp_cell.text = self.company_name
            # 检测因素
            item_cell = info_table.cell(3, 1)
            item_cell.text = str(factor)
            # 采样日期
            date_cell = info_table.cell(3, 6)
            # if self.schedule_col == '采样/送样日期':
            #     date_cell.text = schedule.strftime("%Y年%m月%d日")
            for cell in [code_cell, comp_cell, item_cell, date_cell]:
                p = cell.paragraphs[0]
                p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # type: ignore
                if len(cell.text) >= 14:
                    p.runs[0].font.size = Pt(8)
                else:
                    pass
                    # p.runs[0].font.size = Pt(9)
            # 页脚信息
            core_properties = doc_copy.core_properties
            core_properties.keywords = factor
            # if self.schedule_col == '采样/送样日期':
            #     core_properties.comments  = schedule.strftime(r"%Y/%m/%d")
            # 保存到桌面文件夹里
            file_name = f'D{day_i}-个体-{factor}'
            safe_file_name: str = self.convert_safe_filename(file_name)
            file_output_path: str = os.path.join(
                self.output_path,
                safe_file_name
            )
            doc_copy.save(f'{file_output_path}.docx')

    def write_personnel_noise(self, doc: Any) -> None:
        '''将个体噪声信息写入模板'''
        current_factor_info: Dict[str, Any] = self.templates_info['噪声个体']
        # 获得个体噪声信息
        current_factor_df: DataFrame = (
            self.personnel_info_df
            .query('检测因素 == "噪声"')
            .sort_values('采样点编号')
            .reset_index(drop=True)
        )
        # 读取个体噪声模板
        document = deepcopy(doc)
        # 判断需要的记录表的页数
        table_pages: int = (
            math.ceil(
                (len(current_factor_df) - current_factor_info['first_page_rows'])
                / current_factor_info['late_page_rows']
            )
            + 1
        )
        # 根据不同页数，增减表格
        if table_pages == 1:
            # 删除第二页的表格
            rm_table = document.tables[2]
            t = rm_table._element
            t.getparent().remove(t)
            # 删除最后一个段落
            paragraphs = document.paragraphs
            rm_paragraphs1 = paragraphs[-1]
            rm_p1 = rm_paragraphs1._element
            rm_p1.getparent().remove(rm_p1)
            # 删除倒数第二个段落，即模板的第一页的换页符
            rm_paragraphs2 = paragraphs[-2]
            rm_p2 = rm_paragraphs2._element
            rm_p2.getparent().remove(rm_p2)
        elif table_pages == 2:
            pass # 跳过
        else:
            # 循环增加表格
            for _ in range(table_pages - 2):
                # 复制第二页的表格
                cp_table = document.tables[2]
                new_table = deepcopy(cp_table)
                # 在模板末尾增加段落
                new_paragraph = document.add_page_break()
                # 增加复制的表格
                new_paragraph._p.addnext(new_table._element)
                # 再增加一个段落
                document.add_paragraph()
        # 写入信息
        # 处理后的模板的所有表格
        tables = document.tables
        # 分析不同表格的写入信息
        for table_page in range(table_pages):
            # 获得当前表格的相应信息的索引
            if table_page == 0:
                index_first: int = 0
                index_last: int = current_factor_info['first_page_rows'] - 1
            else:
                index_first: int = (
                    current_factor_info['late_page_rows']
                    * (table_page - 1)
                    + current_factor_info['first_page_rows']
                )
                index_last: int = (
                    current_factor_info['first_page_rows']
                    + table_page
                    * current_factor_info['late_page_rows']
                    - 1
                )
            # 筛选出当前表格的信息
            if index_first == index_last:
                current_df: DataFrame = (
                    current_factor_df
                    .query(f'index == {index_first}')
                    .reset_index(drop=True)
                )
            else:
                current_table = tables[table_page + 1]
                current_df: DataFrame = (
                    current_factor_df
                    .query(f'index >= {index_first} and index <= {index_last}')
                    .reset_index(drop=True)
                )
            current_table = tables[table_page + 1]
            # 按行循环选取单元格
            for r_i in range(current_df.shape[0]):
                current_row_list = [
                    current_df.loc[r_i, '采样点编号'],
                    f"{current_df.loc[r_i, '单元']} {current_df.loc[r_i, '工种']}\n",
                    current_df.loc[r_i, '日接触时间'],
                ]
                # 再循环列选取单元格，并写入相应信息
                for i, c_i in enumerate(current_factor_info['available_cols']):
                    current_cell = (
                        current_table.rows[
                            r_i * current_factor_info['item_rows']
                            + current_factor_info['title_rows']
                            ]
                        .cells[c_i]
                    )
                    current_cell.text = str(current_row_list[i])
                    current_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    current_cell.paragraphs[0].runs[0].font.size = Pt(6.5)
                    # current_cell.paragraphs[0].runs[0].font.name = '宋体'
        info_table = tables[0]
        code_cell = info_table.cell(0, 1)
        comp_cell = info_table.cell(1, 1)

        code_cell.text = self.project_number
        code_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
        comp_cell.text = self.company_name
        comp_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
        # [x] 单元格样式
        file_name: str = '个体噪声记录表'
        output_file_path: str = os.path.join(
            self.output_path,
            f'{file_name}.docx'
        )
        document.save(output_file_path)

    def write_direct_reading_factors_docx(self, other_point_factor: str) -> None:
        '''将仪器只读信息写入模板'''
        # [x] 去除重复的检测信息
        # 获得检测因素的信息
        factor_key: str = f'{other_point_factor}定点'
        current_factor_info: Dict[str, Any] = self.templates_info[factor_key]
        join_char: str = current_factor_info['join_char']
        # 获得检测因素的点位信息
        current_factor_df: DataFrame = (
            self.point_info_df
            .query(f'检测因素 == "{other_point_factor}"')
            .sort_values('采样点编号')
            .reset_index(drop=True)
        )
        # 读取检测因素模板
        current_factor_template: str = (
            self.templates_info
            [factor_key]['template_path']
        )
        document = Document(current_factor_template)
        # 判断需要的记录表的页数
        table_pages: int = (
            math.ceil(
                (len(current_factor_df) - current_factor_info['first_page_rows'])
                / current_factor_info['late_page_rows']
            )
            + 1
        )
        # 根据不同页数，增减表格
        if table_pages == 1:
            # 删除第二页的表格
            rm_table = document.tables[2]
            t = rm_table._element
            t.getparent().remove(t)
            # 删除最后一个段落
            paragraphs = document.paragraphs
            rm_paragraphs1 = paragraphs[-1]
            rm_p1 = rm_paragraphs1._element
            rm_p1.getparent().remove(rm_p1)
            # 删除倒数第二个段落，即模板的第一页的换页符
            rm_paragraphs2 = paragraphs[-2]
            rm_p2 = rm_paragraphs2._element
            rm_p2.getparent().remove(rm_p2)
        elif table_pages == 2:
            pass # 跳过
        else:
            # 循环增加表格
            for _ in range(table_pages - 2):
                # 复制第二页的表格
                cp_table = document.tables[2]
                new_table = deepcopy(cp_table)
                # 在模板末尾增加段落
                new_paragraph = document.add_page_break()
                # 增加复制的表格
                new_paragraph._p.addnext(new_table._element)
                # 再增加一个段落
                document.add_paragraph()
        # [x] 写入信息
        # # 处理后的模板的所有表格
        tables = document.tables
        # 分析不同表格的写入信息
        for table_page in range(table_pages):
            # 获得当前表格的相应信息的索引
            if table_page == 0:
                index_first: int = 0
                index_last: int = current_factor_info['first_page_rows'] - 1
            else:
                index_first: int = (
                    current_factor_info['late_page_rows']
                    * (table_page - 1)
                    + current_factor_info['first_page_rows']
                )
                index_last: int = (
                    current_factor_info['first_page_rows']
                    + table_page
                    * current_factor_info['late_page_rows']
                    - 1
                )
            # 筛选出当前表格的信息
            if index_first == index_last:
                current_df: DataFrame = (
                    current_factor_df
                    .query(f'index == {index_first}')
                    .reset_index(drop=True)
                )
            else:
                current_table = tables[table_page + 1]
                current_df: DataFrame = (
                    current_factor_df
                    .query(f'index >= {index_first} and index <= {index_last}')
                    .reset_index(drop=True)
                )
            current_table = tables[table_page + 1]
            # 按行循环选取单元格
            for r_i in range(current_df.shape[0]):
                current_row_list = [
                    current_df.loc[r_i, '采样点编号'],
                    f"{current_df.loc[r_i, '单元']}{join_char}{current_df.loc[r_i, '检测地点']}",
                    current_df.loc[r_i, '日接触时间'],
                ]
                # 再循环列选取单元格，并写入相应信息
                for i, c_i in enumerate(current_factor_info['available_cols']):
                    current_cell = (
                        current_table.rows[
                            r_i * current_factor_info['item_rows']
                            + current_factor_info['title_rows']
                            ]
                        .cells[c_i]
                    )
                    current_cell.text = str(current_row_list[i])
                    current_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    current_cell.paragraphs[0].runs[0].font.size = Pt(9)
                    # current_cell.paragraphs[0].runs[0].font.name = '宋体'
        # [x] 样式调整
        # 写入基本信息
        info_table = tables[0]
        code_cell = (
            info_table
            .rows[current_factor_info['project_num_row']]
            .cells[current_factor_info['project_num_col']]
        )
        comp_cell = (
            info_table
            .rows[current_factor_info['company_name_row']]
            .cells[current_factor_info['company_name_col']]
        )
        code_cell.text = self.project_number
        code_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
        comp_cell.text = self.company_name
        comp_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
        # [x] 样式调整
        # 保存
        file_name: str = f'{other_point_factor}记录表'
        safe_file_name: str = self.convert_safe_filename(file_name)
        output_file_path: str = os.path.join(
            self.output_path, f'{safe_file_name}.docx')
        document.save(output_file_path)

    def write_traveler_docx(self, doc, schedule_day: int) -> None:
        '''将流转单信息写入模板'''
        # 流转单信息
        counted_df = self.output_deleterious_substance_info_dict[f'{schedule_day}']['样品统计']
        # 写入基本信息
        # traveler_path: str = './templates/样品流转单.docx'
        traveler_document = deepcopy(doc)
        project_num_cell = traveler_document.tables[0].rows[0].cells[1]
        project_num_cell.text = self.project_number
        # [x] 样式
        # 判断需要的流转单的页数
        table_pages: int = math.ceil(len(counted_df) / 8)
        for _ in range(table_pages - 1):
            cp_table = traveler_document.tables[0]
            new_table = deepcopy(cp_table)
            cp_paragraph = traveler_document.paragraphs[0]
            last_paragraph = traveler_document.add_page_break()
            last_paragraph._p.addnext(new_table._element)
            traveler_document.add_paragraph(cp_paragraph.text)

        tables = traveler_document.tables

        for table_page in range(table_pages):
            first_index: int = 8 * table_page
            last_index: int = 8 * table_page + 7
            # .reset_index(drop=True)
            current_df: DataFrame = (
                counted_df
                .iloc[first_index : last_index + 1]
                .reset_index()
            )
            current_table = tables[table_page]
            for r_i in range(len(current_df)):
                # current_index_name = current_df.iloc[r_i].name
                # print(current_index_name)
                current_row_list = [
                    current_df.loc[r_i, "编号范围"],  # type: ignore
                    current_df.loc[r_i, "检测因素"],  # type: ignore
                    current_df.loc[r_i, "保存时间"],  # type: ignore
                    current_df.loc[r_i, "总计"],  # type: ignore
                ]
                for c_i in list(range(4)):
                    match_cols_list: List[int] = [0, 1, 3, 4]
                    current_cell = (
                        current_table
                        .rows[r_i + 2]
                        .cells[match_cols_list[c_i]]
                    )
                    current_cell.text = str(current_row_list[c_i])
                    current_cell.paragraphs[0].runs[0].font.size = Pt(7.5)
                    current_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    # [x] 单元格样式
                    if '\\n' in current_cell.text:
                        new_text: str = current_cell.text.replace('\\n', '\n')
                        current_cell.text = new_text
                        current_cell.paragraphs[0].runs[0].font.size = Pt(7.5)
                        current_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
        file_name: str = f'D{schedule_day}-样品流转单'
        safe_file_name: str = re.sub(r'[?*/\<>:"|]', ',', file_name)
        # output_path = f'{os.path.expanduser("~/Desktop")}/{self.project_number}记录表'
        # if not os.path.exists(self.output_path):
        #     os.mkdir(self.output_path)
        # else:
        #     pass
        output_file_path: str = os.path.join(
            self.output_path, f'{safe_file_name}.docx')
        traveler_document.save(output_file_path)



    def get_counted_df_save_info(self, factor: str) -> str:
        '''获得样品统计df里的各个检测因素的保存时间'''
        if factor.count('|') == 0:
            first_factor: str = factor
        else:
            first_factor: str = factor.split('|')[0]

        if first_factor in self.factor_reference_df['标识检测因素'].values:
            save_info_df: DataFrame = (
                self.factor_reference_df
                .query("标识检测因素 == @first_factor")
                .reset_index(drop=True)
            )
            save_info: str = str(save_info_df.loc[0, '保存时间'])
        else:
            save_info: str = '/'
        return save_info

    # [x] 更新已占用编号数量

    def refresh_engaged_num(
        self,
        current_df: DataFrame,
        current_type: str,
        engaged_num: int
    ) -> int:
        '''更新已占用样品编号数'''
        # 按照df类型来更新编号
        # [x] 如果df长度为0时要
        # [x] 更新，使用字典模式
        type_num_dict: Dict[str, str] = {
            '空白': '空白编号',
            '定点': '终止编号',
            '个体': '个体编号',
        }
        if current_df.empty == False and current_type in self.default_types_order:
            new_engaged_num: int = (
                current_df[type_num_dict[current_type]]
                .astype(int)
                .max()
            )
            return new_engaged_num
        else:
            return engaged_num

    # [x] 列表的自定义排序

    def custom_sort(self, str_list: List[str], key_list: List[str]) -> List[str]:
        '''
        列表的自定义排序
        '''
        if str_list[0] in key_list:
            sorted_str_list: List[str] = sorted(
                str_list,
                key=lambda x: key_list.index(x)
            )
            return sorted_str_list
        else:
            return str_list
    # [x] 获得空白、定点和个体的编号范围

    def get_blank_count_range(self, blank_df: DataFrame) -> str:
        '''获得空白的编号范围'''
        if blank_df['空白数量'] != 0:
            blank_str: str = (
                f'{blank_df["空白编号"]:0>4d}-1, {blank_df["空白编号"]:0>4d}-2'
            )
            return blank_str
        else:
            return ''

    def get_point_count_range(self, point_df: DataFrame) -> str:
        '''获得定点的编号范围'''
        if point_df['定点数量'] == 0:
            return ''
        elif point_df['定点数量'] == 1:
            point_str: str = f'{point_df["起始编号"]:0>4d}'
            return point_str
        else:
            point_str: str = (
                f'{point_df["起始编号"]:0>4d}-{point_df["终止编号"]:0>4d}'
            )
            return point_str

    def get_personnel_count_range(self, personnel_df: DataFrame) -> str:
        '''获得个体的编号范围'''
        if personnel_df['个体数量'] == 0:
            return ''
        elif personnel_df['个体数量'] == 1:
            personnel_str: str = f'{personnel_df["个体起始编号"]:0>4d}'
            return personnel_str
        else:
            personnel_str: str = (
                f'{personnel_df["个体起始编号"]:0>4d}-{personnel_df["个体终止编号"]:0>4d}'
            )
            return personnel_str

    # [x] 将编号范围转换为字符串
    def get_range_str(self, counted_df: DataFrame):
        '''将编号范围转换为字符串'''
        range_list = [
            counted_df['空白编号范围'],
            counted_df['定点编号范围'],
            counted_df['个体编号范围']
        ]
        range_list = [i for i in range_list if i != '']
        range_str = ', '.join(range_list)  # type: ignore
        return range_str

    # [x] 从定点和个体的采样日程获得项目的总日程

    def get_schedule_days(self) -> int:
        '''从定点和个体的采样日程获得项目的总日程'''
        point_schedule_days: int = self.point_info_df['采样日程'].max()
        personnel_schedule_days: int = self.personnel_info_df['采样日程'].max()
        schedule_days: int = max(
            point_schedule_days,
            personnel_schedule_days
        )
        return schedule_days

    # [x] 转换文件名为可保存的文件名
    def convert_safe_filename(self, file_name: str) -> str:
        '''转换文件名为可保存的文件名'''
        safe_file_name: str = re.sub(r'[?*/\<>:"|]', ',', file_name)
        return safe_file_name

    # [x] 添加采样时间排序功能
    def time_manage(self, schedule_day: int):
        '''添加采样时间排序功能'''
        # current_point_df: DataFrame = self.output_deleterious_substance_info_dict[f'{schedule_day}']['定点']
        pass
