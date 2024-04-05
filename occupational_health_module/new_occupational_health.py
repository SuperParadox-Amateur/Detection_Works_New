'''处理系统生成编号，并写入模板里'''
# import os
from copy import deepcopy
import json
import math
from pathlib import Path
from decimal import ROUND_HALF_DOWN, Decimal
import re
from typing import Any, Dict, List, Optional

from nptyping import DataFrame
import pandas as pd
from pandas import CategoricalDtype
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


class NewOccupationalHealthItemInfo():
    '''系统生成的职业卫生样品编号处理信息'''
    def __init__(
            self,
            project_number: str,
            company_name: str,
            raw_df: DataFrame[Any],
            in_templates_categories: Optional[List[str]],
            is_all_factors_split: bool = False,
            ) -> None:
        all_templates_categories: List[str] = [
            '定点有害物质',
            '个体有害物质',
            '个体噪声',
            '仪器直读因素',
            '样品流转单',
        ]
        self.company_name: str = company_name  # 公司名称
        self.project_number: str = project_number  # 项目编号
        # self.templates_path_dict: Dict[str, Any] = (
        #     self.read_json_to_dict('./info_files/默认模板路径信息.json')
        # )
        self.templates_info_dict: Dict[str, Any] = (
            self.read_json_to_dict('./info_files/默认模板格式信息.json')
        )
        self.is_all_factors_split: bool = is_all_factors_split # 是否所有检测因素按照日期分开
        self.factor_reference_df: DataFrame[Any] = self.get_occupational_health_factor_reference()
        self.df: DataFrame[Any] = self.initialize_df(raw_df)
        self.schedule_col: str = self.initialize_schedule()
        self.schedule_list: List[Any] = self.get_schedule_list()
        # self.templates_categories: List[str] = self.initialize_in_templates_categories(in_templates_categories)
        self.templates_categories: List[str] = (
            all_templates_categories
            if in_templates_categories is None
            else in_templates_categories
        )
        self.blank_df: DataFrame[Any] = self.initialize_blank_df()
        self.point_df: DataFrame[Any] = self.initialize_point_df()
        self.personnel_df: DataFrame[Any] = self.initialize_personnel_df()
        self.stat_df: DataFrame[Any] = self.initialize_stat_df()
        # self.output_path: str = os.path.join(
        #     os.path.expanduser("~/Desktop"),
        #     f'{self.project_number}记录表'
        # )
        self.output_path: Path = Path.home()/'Desktop'/f'{project_number}记录表'


# 初始化

    # 默认获得职业卫生所有检测因素的参考信息
    def get_occupational_health_factor_reference(self) -> DataFrame[Any]:
        '''
        获得职业卫生所有检测因素的参考信息
        '''
        # reference_path: str = os.path.join(
        #     'info_files/检测因素参考信息.csv'
        # )
        reference_path: Path = Path('info_files/检测因素参考信息.csv')
        reference_df: DataFrame[Any] = pd.read_csv(reference_path)  # type: ignore
        # 增加不同列的空值为不同的数值
        fill_dict: Dict[str, Any] = {
            '采样仪器': '/',
            '保存时间': '/',
            '流量*时间': '/',
            '定点采样流速': 0.0,
            '定点采样时间': 0,
            '个体采样流速': 0.0,
            '个体采样时间': 0,
        }
        reference_df: DataFrame[Any] = reference_df.fillna(fill_dict)  # type: ignore
        return reference_df

    def read_json_to_dict(self, json_file: str) -> Dict[str, Any]:
        '''读取json文件并转换为dict'''
        with open(json_file, 'r', encoding='utf-8') as f:
            string: str = f.read()
            jdict: Dict[str, Any] = json.loads(string)
        return jdict

    def initialize_df(self, raw_df: DataFrame[Any]) -> DataFrame[Any]:
        '''初始化所有样品信息'''
        available_cols: List[str] = [
            '样品类型',
            '样品编号',
            '样品名称',
            '检测参数',
            '采样/送样日期',
            '样品描述',
            '单元',
            '工种/岗位',
            '检测地点',
            '测点编号',
            '第几天',
            '第几个频次',
            '采样方式',
            '作业人数',
            '日接触时长/h',
            '周工作天数/d',
        ]
        fillna_dict: Dict[str, Any] = {
            '日接触时长/h': 0.0,
            '周工作天数/d': 0.0,
        }
        # cols_dtypes = {
            # '样品类型': str,
            # '样品编号': str,
            # '样品名称': str,
            # '检测参数': str,
            # '采样/送样日期': 'datetime64[ns]',
            # '单元': str,
            # '工种/岗位': str,
            # '检测地点': str,
            # '测点编号': int,
            # '第几天': int,
            # '第几个频次': int,
            # '采样方式': str,
            # '作业人数': int,
            # '日接触时长/h': float,
            # '周工作天数/d': float,
        # }
        df: DataFrame[Any] = raw_df[available_cols].sort_values(by='测点编号') # type: ignore
        df: DataFrame[Any] = df.fillna(fillna_dict) # type: ignore
        # df: DataFrame[Any] = df.astype(cols_dtypes) # type: ignore
        #  将检测参数列转换为category类型用于排序
        # （取消，因为会导致groupby性能下降，甚至失败）
        # factor_list: List[str] = df['检测参数'].unique().tolist()
        # sorted_factor_list: List[str] = sorted(factor_list, key=lambda x: x.encode('gbk'))
        # factor_order = CategoricalDtype(sorted_factor_list, ordered=True)
        # df['检测参数'] = df['检测参数'].astype(factor_order)
        # df['样品编号'] = df['样品编号'].replace(self.project_number, '', inplace=True) # type: ignore
        df['样品编号'] = df['样品编号'].apply(self.parse_sample_num) # type: ignore
        # [ ] 如果采样日程的相关列都是空，则填充0
        # df['样品编号'] = (
        #     df['样品编号']
        #     .apply(self.handle_num_str)
        #     .reset_index(drop=True)
        # ) # type: ignore

        return df

    def initialize_blank_df(self) -> DataFrame[Any]:
        '''初始化空白信息'''
        raw_blank_df: DataFrame[Any] = (
            self # type: ignore
            .df
            .query('样品类型 == "空白样"')
            .reset_index(drop=True)
        )
        blank_df: DataFrame[Any] = (
            raw_blank_df # type: ignore
            .pivot(
                index=['检测参数', self.schedule_col],
                columns='第几个频次',
                values='样品编号'
            )
            .rename(columns={1: '空白编号1', 2: '空白编号2'})
            .reset_index(drop=False)
        )

        return blank_df

    def initialize_point_df(self) -> DataFrame[Any]:
        '''初始化定点信息'''
        query_str: str = (
            '样品类型 == "普通样"'
            ' and '
            '采样方式 == "定点"'
            ' and '
            '样品名称 != "工作场所物理因素"'
            ' and '
            '样品描述 != "仪器直读"'
            ' and '
            '样品编号 != "/"'
        )
        raw_point_df: DataFrame[Any] = (
            self # type: ignore
            .df
            .query(query_str)
            .reset_index(drop=True)
        )
        if raw_point_df.shape[0] != 0:
            raw_point_df['样品编号'] = (
                raw_point_df['样品编号'] # type: ignore
                .astype(int)
            )
            groupby_point_df: DataFrame[Any] = (
                raw_point_df # type: ignore
                .groupby(
                    [
                        '测点编号',
                        '单元',
                        '检测地点',
                        '工种/岗位',
                        '检测参数',
                        self.schedule_col,
                        r'日接触时长/h'
                    ]
            )
            ['样品编号']
            .agg(list)
            .reset_index(drop=False)
            )

            groupby_point_df['采样数量/天'] = (
                groupby_point_df # type: ignore
                ['样品编号']
                .apply(len)
            )
            # [x] 是否合并代表时长列要改进
            groupby_point_df['是否合并代表时长'] = (
                groupby_point_df # type: ignore
                .apply(
                    lambda df: True if df[r'日接触时长/h']/df['采样数量/天']<0.25 else False, # type: ignore
                    axis=1
                )
            )
            point_df: DataFrame[Any] = groupby_point_df.merge( # type: ignore
                self.blank_df,
                on=['检测参数', self.schedule_col],
                how='left'
            )
            point_df['代表时长'] = (
                point_df # type: ignore
                .apply(
                    lambda df: self.get_exploded_contact_duration( # type: ignore
                        df[r'日接触时长/h'], df[r'采样数量/天']#, 4 # type: ignore
                    ),
                    axis=1
                )
            )
            point_df['空白编号1'] = point_df['空白编号1'].fillna('-') # type: ignore
            point_df['空白编号2'] = point_df['空白编号2'].fillna('-') # type: ignore
        else:
            empty_cols: List[str] = [
                '测点编号',
                '单元',
                '检测地点',
                '工种/岗位',
                '检测参数',
                '采样/送样日期',
                '日接触时长/h',
                '样品编号',
                '采样数量/天',
                '是否合并代表时长',
                '空白编号1',
                '空白编号2',
                '代表时长'
            ]
            point_df = pd.DataFrame(columns=empty_cols)

        return point_df

    def initialize_personnel_df(self) -> DataFrame[Any]:
        '''初始化个体信息'''
        query_str: str = (
            '样品类型 == "普通样"'
            ' and '
            '采样方式 == "个体"'
            ' and '
            '样品名称 != "工作场所物理因素"'
            ' and '
            '样品编号 != "/"'
        )
        personnel_df: DataFrame[Any] = (
            self # type: ignore
            .df
            .query(query_str)
            .reset_index(drop=True)
        )
        # 去除空行
        new_personnel_df: DataFrame[Any] = (
            personnel_df # type: ignore
            .dropna(how='all')
            .reset_index(drop=True)
        )
        if not new_personnel_df.empty:
            new_personnel_df['样品编号'] = (
                new_personnel_df['样品编号'] # type: ignore
                .astype(int)
            )
        return new_personnel_df

    def initialize_stat_df(self) -> DataFrame[Any]:
        '''获得样本统计信息'''
        # 筛选出定点和个体的所有信息
        query_str: str = (
            '样品名称 != "工作场所物理因素"'
            ' and '
            '样品类型 == "普通样"'
            ' and '
            '样品编号 != "/"'
        )
        df: DataFrame[Any] = (
            self.df # type: ignore
            .query(query_str)
            .reset_index()
        )
        if df.shape[0] != 0:
            df['样品编号'] = df['样品编号'].astype(int) # 转为整数型 # type: ignore
            # 按照日程和检测参数分组并转换为列表
            groupby_df: DataFrame[Any] = (
                df # type: ignore
                .groupby([self.schedule_col, '检测参数'])
                ['样品编号']
                .agg(list)
                # .agg(self.convert_merge_range)
                .reset_index()
            )
            # 定点和个体的数量
            groupby_df['样品数量'] = groupby_df['样品编号'].apply(len) # type: ignore
            groupby_df['样品编号'] = groupby_df['样品编号'].apply(self.convert_merge_range) # type: ignore
            # 合并空白
            merged_df: DataFrame[Any] = groupby_df.merge( # type: ignore
                self.blank_df, # type: ignore
                on=['检测参数', '采样/送样日期'],
                how='left'
            )
            # 将定点和个体的编号与空白编号合并为一个列表
            merged_df['空白编号1'] = (
                merged_df['空白编号1'] # type: ignore
                .fillna('-')
            )
            merged_df['空白编号2'] = (
                merged_df['空白编号2'] # type: ignore
                .fillna('-')
            )
            # 空白样品数量
            merged_df['空白数量'] = (
                merged_df['空白编号1'] # type: ignore
                .apply(lambda x: 2 if x != '-' else 0) # type: ignore
            )
            merged_df['空白编号'] = (
                merged_df # type: ignore
                .apply(
                    lambda x: [x['空白编号1'], x['空白编号2']], # type: ignore
                    axis=1
                )
            )
            merged_df['全部样品编号'] = (
                merged_df # type: ignore
                .apply(
                    lambda x: x['空白编号'] + x['样品编号'], # type: ignore
                    axis=1
                )
            )
            # 去除无空白
            merged_df['全部样品编号'] = (
                merged_df['全部样品编号'] # type: ignore
                .apply(lambda x: [i for i in x if i != '-']) # type: ignore
            )
            # 所有样品数量
            merged_df['全部样品数量'] = (
                merged_df['空白数量']
                + merged_df['样品数量']
            )
            # 检测参数按照拼音排序
            factor_list: List[str] = (
                merged_df['检测参数'] # type: ignore
                .unique()
                .tolist()
            )
            sorted_factor_list: List[str] = sorted(
                factor_list,
                key=lambda x: x.encode('gbk')
            )
            factor_order: CategoricalDtype = CategoricalDtype(sorted_factor_list, ordered=True)
            merged_df['检测参数'] = merged_df['检测参数'].astype(factor_order) # type: ignore
            merged_df: DataFrame[Any] = (
                merged_df # type: ignore
                .sort_values(
                    by=[self.schedule_col, '检测参数'],
                    ascending=True,
                    ignore_index=True
                )
            )
            # 加上检测参数的保存时间
            merged_df['标识检测参数'] = merged_df['检测参数'].apply(self.get_split_str_first) # type: ignore
            all_merged_df: DataFrame[Any] = (
                merged_df # type: ignore
                .merge(
                    self.factor_reference_df,
                    left_on='标识检测参数',
                    right_on='标识检测因素',
                    how='left'
                )
                .fillna({'保存时间': '/'})
            )
        else:
            all_merged_cols: List[str] = [
                '采样/送样日期',
                '检测参数',
                '样品编号',
                '样品数量',
                '空白编号1',
                '空白编号2',
                '空白数量',
                '空白编号',
                '全部样品编号',
                '全部样品数量',
                '标识检测参数',
                '标识检测因素',
                '样品收集器',
                '采样仪器',
                '是否仪器直读',
                '收集方式',
                '是否需要空白',
                '保存时间',
                '流量*时间',
                '定点采样流速',
                '定点采样时间',
                '个体采样流速',
                '个体采样时间',
                '备注',
                '复合因素代码'
            ]
            all_merged_df: DataFrame[Any] = pd.DataFrame(columns=all_merged_cols)

        return all_merged_df

    def initialize_schedule(self) -> str:
        '''初始化采样日程'''
        if self.df['采样/送样日期'].isnull().all(): # type: ignore
            schedule_col: str = '第几天'
        else:
            schedule_col: str = '采样/送样日期'
        return schedule_col

    def get_schedule_list(self) -> List[Any]:
        '''获得采样日程'''
        if self.schedule_col == '采样/送样日期':
            self.df[self.schedule_col] = pd.to_datetime(self.df[self.schedule_col]) # type: ignore
        # 可能是整数或者是日期
        schedule_list: List[Any] = (
            self.df[self.schedule_col] # type: ignore
            .drop_duplicates()
            .tolist()
        )
        # if self.schedule_col == '采样/送样日期':
        #     schedule_list = [
        #         datetime.strptime(i, '%Y-%m-%d').date() for i in schedule_list # type: ignore
        #     ]
        return sorted(schedule_list)

    # def initialize_in_templates_categories(self, in_templates_categories: Optional[List[str]]) -> List[str]:
    #     '''初始化要处理的模板分类'''
    #     if in_templates_categories == None:
    #         templates_categories: List[str] = [
    #             '定点有害物质',
    #             '个体有害物质',
    #             '个体噪声',
    #             '仪器直读因素',
    #             '样品流转单',
    #         ]
    #     else:
    #         templates_categories: List[str] = in_templates_categories
    #     return templates_categories

# 自定义函数

    def parse_sample_num(self, sample_num: str) -> str:
        '''整理样品编号'''
        if sample_num != '/':
            parsed_sample_num: str = sample_num.replace(self.project_number, '')
            return parsed_sample_num
        else:
            return '/'


    def get_exploded_contact_duration(
        self,
        duration: float,
        size: int = 4
    ) -> List[float]:
        '''获得分开的接触时间，使用十进制来计算'''
        # 接触时间和数量转为十进制
        time_dec: Decimal = Decimal(str(duration))
        size_dec: Decimal = Decimal(str(size))
        time_list_dec: List[Decimal] = [] # 存放代表时长列表
        # 判断接触时间的小数位数
        if duration == int(duration):
            time_prec: int = 0
        else:
            time_prec: int = int(time_dec.as_tuple().exponent)
        # 确定基本平均值的小数位数
        time_prec_dec_dict: Dict[int, Decimal] = {
            0: Decimal('0'),
            -1: Decimal('0.0'),
            -2: Decimal('0.0')
        }
        prec_dec_str: Decimal = time_prec_dec_dict[time_prec]
        # 如果接触时间不能让每个代表时长大于0.25，则不分开
        if time_dec < Decimal('0.25') * size_dec:
            time_list_dec.append(time_dec)
        elif time_dec < Decimal('0.5') * size_dec:
            front_time_list_dec: List[Decimal] = [
                Decimal('0.25')] * (int(size) - 1)
            last_time_dec: Decimal = time_dec - sum(front_time_list_dec)
            time_list_dec.extend(front_time_list_dec)
            time_list_dec.append(last_time_dec)
        elif time_dec < Decimal('0.7') * size_dec:
            front_time_list_dec: List[Decimal] = [
                Decimal('0.5')] * (int(size) - 1)
        else:
            judge_result: Decimal = time_dec / size_dec
            for _ in range(int(size) - 1):
                result: Decimal = judge_result.quantize(prec_dec_str, ROUND_HALF_DOWN)
                time_list_dec.append(result)
            last_result: Decimal = time_dec - sum(time_list_dec)
            time_list_dec.append(last_result)
        time_list: List[float] = list(map(float, time_list_dec))
        return time_list

    def convert_merge_range(self, raw_lst: List[int]) -> List[str]:
        '''将编号列表里连续的编号合并，并转换为列表'''
        lst: List[int] = sorted(raw_lst)
        # lst: List[int] = [1, 2, 3, 4, 5, 7, 8, 9, 10, 11, 13, 14, 15, 17, 18]
        all_range_list: List[List[int]] = []
        current_range: List[int] = []
        lst.extend([0])

        for i, num in enumerate(lst[:-1]):
            start: int = num
            current_range.append(start)
            end: int = num + 1
            if end == lst[i + 1]:
                # range.append(start)
                pass
            else:
                all_range_list.append(current_range)
                current_range = []

        range_str_list: List[str] = []
        for range_list in all_range_list:
            if len(range_list) != 1:
                range_str: str = f'{range_list[0]:>04d}--{range_list[-1]:>04d}'
                range_str_list.append(range_str)
            else:
                range_str: str = f'{range_list[0]:>04d}'
                range_str_list.append(range_str)

        return range_str_list

    def get_split_str_first(self, input_str: str) -> str:
        '''获得检测因素的标识检测因素'''
        str_list: List[str] = input_str.split('|')
        return str_list[0]

    # def get_template_abs_path(self, templates_path_dict: Dict[str, str]) -> Dict[str, str]:
    #     '''获得模板的绝对路径'''
    #     templates_path_abs_dict: Dict[str, str] = {}
    #     for i, j in templates_path_dict.items():
    #         abs_path: str = os.path.join(
    #             os.path.abspath(os.path.join(os.getcwd(), "..")),
    #             j
    #         )
    #         templates_path_abs_dict[i] = abs_path
    #     return templates_path_abs_dict

    def is_writable_to_templates(self, target_df: DataFrame[Any], template_category: str) -> bool:
        '''判断当前检测因素的dataframe是否执行写入'''
        is_not_len_zero: bool = len(target_df) != 0
        is_planed: bool = True if template_category in self.templates_categories else False
        is_writable: bool = is_not_len_zero and is_planed
        return is_writable

# 写入模板文件中

    def write_to_templates(self) -> None:
        '''将全部信息写入全部对应模板'''
        # [ ] 增加向剩余的模板写入日期功能
        # 桌面保存文件夹不存在则创建
        if not self.output_path.exists():
            Path.mkdir(self.output_path)

        # 按照日程写入对应模板
        # [x] 有害物质要分日程和不同检测因素
        for day_idx, schedule in enumerate(self.schedule_list):
            # 定点有害物质
            current_point_df: DataFrame[Any] = (
                self.point_df  # type:ignore
                [self.point_df[self.schedule_col] == schedule]
                .reset_index(drop=True)
            )
            # 当天的定点有害物质数量不是0，并且在计划里就写入
            is_writable_point: bool = self.is_writable_to_templates(current_point_df, '定点有害物质')
            if is_writable_point:
                doc_point: Any = Document(self.templates_info_dict['有害物质定点']['template_path'])
                point_info_dict: Dict[str, Any] = self.templates_info_dict['有害物质定点']
                factors: List[str] = current_point_df['检测参数'].drop_duplicates().tolist()
                # 获得当前检测因素的dataframe
                for current_factor in factors:
                    # 导入模板
                    doc_point_copy: Any = deepcopy(doc_point)
                    current_factor_df: DataFrame[Any] = (
                        current_point_df[current_point_df['检测参数'] == current_factor]
                        .sort_values(by='测点编号')
                        .reset_index(drop=True)
                    )
                    self.write_point_deleterious_substance(
                        doc_point_copy,
                        point_info_dict,
                        current_factor_df, # type: ignore
                        current_factor, # type: ignore
                        day_idx,
                        schedule
                    )
            # 个体有害物质
            current_personnel_df: DataFrame[Any] = ( # type: ignore
                self.personnel_df  # type:ignore
                [self.personnel_df[self.schedule_col] == schedule]
                .reset_index(drop=True)
            )
            # 当天的个体有害物质数量不是0，并且在计划里就写入
            is_writable_personnel: bool = self.is_writable_to_templates(current_personnel_df, '个体有害物质')
            if is_writable_personnel:
                doc_personnel: Any = Document(self.templates_info_dict['有害物质个体']['template_path'])
                personnel_info_dict: Dict[str, Any] = self.templates_info_dict['有害物质个体']
                factors: List[str] = current_personnel_df['检测参数'].drop_duplicates().tolist() # type: ignore
                # 获得当前检测因素的dataframe
                for current_factor in factors: # type: ignore
                    # 导入模板
                    doc_personnel_copy: Any = deepcopy(doc_personnel)
                    current_factor_df: DataFrame[Any] = ( # type: ignore
                        current_personnel_df[current_personnel_df['检测参数'] == current_factor] # type: ignore
                        .sort_values(by='测点编号')
                        .reset_index(drop=True)
                    )
                    self.write_personnel_deleterious_substance(
                        doc_personnel_copy,
                        personnel_info_dict,
                        current_factor_df, # type: ignore
                        current_factor, # type: ignore
                        day_idx,
                        schedule
                    )
            # 样品流转单
            current_traveler_df: DataFrame[Any] = ( # type: ignore
                self.stat_df # type: ignore
                [self.stat_df[self.schedule_col] == schedule]
                .reset_index(drop=True)
            )
            # 当天的有害物质数量不是0，并且在计划里就写入
            is_writable_traveler: bool = self.is_writable_to_templates(current_traveler_df, '样品流转单')
            if is_writable_traveler:
                traveler_doc = Document(self.templates_info_dict['流转单']['template_path'])
                self.write_traveler_docx(current_traveler_df, traveler_doc, day_idx, schedule) # type: ignore

        # [x] 直读检测因素设置为是否按照日程分开
        # 仪器直读因素
        other_factors: List[str] = [
            "一氧化碳",
            "噪声",
            "高温",
            "工频电场",
        ]
        if self.is_all_factors_split:
            for day_idx, schedule in enumerate(self.schedule_list):
                # 个体噪声
                # 获得个体噪声信息
                query_str: str = (
                    '采样方式 == "个体"'
                    ' and '
                    '检测参数 == "噪声"'
                    ' and '
                    f'`{self.schedule_col}` == @schedule'
                )
                current_personnel_noise_df: DataFrame[Any] = (
                    self.df # type: ignore
                    .query(query_str)
                    .sort_values('测点编号')
                    .reset_index(drop=True)
                )
                # 当天的个体噪声数量不是0，并且在计划里就写入
                is_writable_personnel_noise: bool = self.is_writable_to_templates(
                    current_personnel_noise_df, '个体噪声'
                )
                if is_writable_personnel_noise:
                    doc_personnel_noise: Any = Document(
                        self.templates_info_dict['噪声个体']['template_path']
                    )
                    personnel_noise_info_dict: Dict[str, Any] = self.templates_info_dict['噪声个体']
                    self.write_personnel_noise(
                        current_factor_df=current_personnel_noise_df,
                        doc=doc_personnel_noise,
                        current_factor_info = personnel_noise_info_dict,
                        day_idx = day_idx,
                        schedule = schedule
                    )
                # 不同直读检测因素调用不同方法处理
                for factor in other_factors:
                    factor_query_str: str = (
                        f'检测参数 == "{factor}"'
                        ' and '
                        '采样方式 == "定点"'
                        ' and '
                        f'`{self.schedule_col}` == @schedule'
                    )
                    factor_df: DataFrame[Any] = (
                        self.df # type: ignore
                        .query(factor_query_str)
                    )
                    # 当天的直读检测因素数量不是0，并且在计划里就写入
                    is_writable_direct_reading_factor: bool = self.is_writable_to_templates(
                        factor_df, '仪器直读因素'
                    )
                    if is_writable_direct_reading_factor:
                        self.write_direct_reading_factors_docx(
                            direct_reading_factor = factor,
                            day_idx=day_idx,
                            schedule=schedule
                        )
        else:
            # 个体噪声
            # 获得个体噪声信息
            query_str: str = (
                '采样方式 == "个体"'
                ' and '
                '检测参数 == "噪声"'
            )
            current_personnel_noise_df: DataFrame[Any] = (
                self.df # type: ignore
                .query(query_str)
                .sort_values('测点编号')
                .reset_index(drop=True)
            )
            # 当天的个体噪声数量不是0，并且在计划里就写入
            is_writable_personnel_noise: bool = self.is_writable_to_templates(
                current_personnel_noise_df, '个体噪声'
            )
            if is_writable_personnel_noise:
                doc_personnel_noise: Any = Document(self.templates_info_dict['噪声个体']['template_path'])
                personnel_noise_info_dict: Dict[str, Any] = self.templates_info_dict['噪声个体']
                self.write_personnel_noise(
                    current_factor_df=current_personnel_noise_df,
                    doc=doc_personnel_noise,
                    current_factor_info = personnel_noise_info_dict,
                    # day_idx = day_idx,
                    # schedule = schedule
                )
            # 直读检测因素
            # 不同直读检测因素调用不同方法处理
            for factor in other_factors:
                factor_query_str: str = (
                    f'检测参数 == "{factor}"'
                    ' and '
                    '采样方式 == "定点"'
                )
                factor_df: DataFrame[Any] = (
                    self.df # type: ignore
                    .query(factor_query_str)
                )
                # 当天的直读检测因素数量不是0，并且在计划里就写入
                is_writable_direct_reading_factor: bool = self.is_writable_to_templates(
                    factor_df, '仪器直读因素'
                )
                if is_writable_direct_reading_factor:
                    self.write_direct_reading_factors_docx(
                        direct_reading_factor = factor,
                        # day_idx=day_idx,
                        # schedule=schedule
                        )

    def write_point_deleterious_substance(
            self,
            doc_copy: Any,
            info_dict: Dict[str, Any],
            current_factor_df: DataFrame[Any],
            factor: str,
            day_idx: int,
            schedule: Any
        ) -> None:
        '''将相应定点有害物质信息写入对应模板'''
        # 计算需要的记录表页数
        table_pages: int = (
            math
            .ceil(
                (len(current_factor_df) - info_dict['first_page_rows'])
                / info_dict['late_page_rows'] + 1
            )
        )

        # 按照页数来增减表格数量
        if table_pages == 1:
            rm_table = doc_copy.tables[2]
            t = rm_table._element
            t.getparent().remove(t)
            rm_page_break = doc_copy.paragraphs[-2]
            pg = rm_page_break._element
            pg.getparent().remove(pg)
            rm_page_break2 = doc_copy.paragraphs[-2]
            pg2 = rm_page_break2._element
            pg2.getparent().remove(pg2)
        elif table_pages == 2:
            pass
        else:
            for _ in range(table_pages - 2):
                cp_table = doc_copy.tables[2]
                new_table = deepcopy(cp_table)
                new_paragraph = doc_copy.add_page_break()
                new_paragraph._p.addnext(new_table._element)
                doc_copy.add_paragraph()
        # 确定不同的表格要填入的信息范围
        tables = doc_copy.tables
        for table_page in range(table_pages):
            if table_page == 0:
                index_first: int = 0
                index_last: int = info_dict['first_page_rows'] - 1
            else:
                index_first: int = (
                    info_dict['late_page_rows']
                    * (table_page - 1)
                    + info_dict['first_page_rows']
                )
                index_last: int = (
                    info_dict['first_page_rows']
                    + table_page
                    * info_dict['late_page_rows']
                    - 1
                )
            current_df: DataFrame[Any] = (
                current_factor_df # type: ignore
                .query(f'index >= {index_first} and index <= {index_last}')
                .reset_index(drop=True)
            )
            # 向指定表格填写数据
            current_table = tables[table_page + 1]
            for r_i in range(current_df.shape[0]):
                row_info = {
                    '采样点编号': current_df.loc[r_i, '测点编号'],
                    '采样岗位': f"{current_df.loc[r_i, '单元']}\n{current_df.loc[r_i, '检测地点']}",
                    '空白编号1': current_df.loc[r_i, '空白编号1'],
                    '空白编号2': current_df.loc[r_i, '空白编号2'],
                    '样品编号': current_df.loc[r_i, '样品编号'],
                    '代表时长': current_df.loc[r_i, '代表时长'],
                    '是否合并代表时长': current_df.loc[r_i, '是否合并代表时长'],
                }
                # 采样点编号单元格
                cell1 = current_table.cell(r_i * 6 + 2, 0)
                cell1.text = str(row_info['采样点编号'])
                cell1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                cell1.paragraphs[0].runs[0].font.size = Pt(8)
                # 采样岗位单元格
                cell2 = current_table.cell(r_i * 6 + 2, 1)
                cell2.text = row_info['采样岗位']
                cell2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                cell2.paragraphs[0].runs[0].font.size = Pt(7.5)
                #[x] 样品编号加上项目编号前缀
                # 空白编号单元格，只写入第一行
                if table_page == 0 and r_i == 0 and row_info['空白编号1'] != '-':
                    cell3_1 = current_table.cell(r_i * 6 + 2, 2)
                    cell3_1.text = f"{self.project_number}{row_info['空白编号1']}"
                    cell3_2 = current_table.cell(r_i * 6 + 3, 2)
                    cell3_2.text = f"{self.project_number}{row_info['空白编号2']}"
                    cell3_1.paragraphs[0].runs[0].font.size = Pt(8)
                    cell3_2.paragraphs[0].runs[0].font.size = Pt(8)
                else:
                    pass
                # 样品编号单元格
                for n_i, num in enumerate(row_info['样品编号']): # type: ignore
                    cell4 = current_table.cell(r_i * 6 + n_i + 4, 2)
                    cell4.text = f"{self.project_number}{num:0>4d}"
                    cell4.paragraphs[0].runs[0].font.size = Pt(8)
                # 代表时长
                for n_i, duration in enumerate(row_info['代表时长']): # type: ignore
                    cell5 = current_table.cell(r_i * 6 + n_i + 4, 9)
                    cell5.text = str(duration)
                    cell5.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    cell5.paragraphs[0].runs[0].font.size = Pt(9)
                # 是否合并代表时长
                if row_info['是否合并代表时长'] == True:
                    merge_len: int = len(row_info['样品编号']) # type: ignore
                    merge_cell1 = current_table.cell(r_i * 6 + 4, 9)
                    merge_cell2 = current_table.cell(r_i * 6 + merge_len + 3, 9)
                    merge_cell1.merge(merge_cell2)
        #[x] 写入项目基本信息
        info_table = tables[0]
        # 项目编号
        code_cell = info_table.cell(0, 1)
        code_cell.text = self.project_number
        # 单位
        comp_cell = info_table.cell(0, 4)
        comp_cell.text = self.company_name
        # 检测因素
        item_cell = info_table.cell(3, 1)
        item_cell.text = str(factor)
        # 采样日期
        date_cell = info_table.cell(3, 6)
        if self.schedule_col == '采样/送样日期':
            date_cell.text = schedule.strftime("%Y年%m月%d日")
        for cell in [code_cell, comp_cell, item_cell, date_cell]:
            p = cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # type: ignore
            if len(cell.text) >= 14:
                p.runs[0].font.size = Pt(8)
            else:
                pass
                # p.runs[0].font.size = Pt(9)
        # 页脚信息
        core_properties = doc_copy.core_properties
        core_properties.keywords = factor
        if self.schedule_col == "采样/送样日期":
            core_properties.comments  = schedule.strftime(r"%Y/%m/%d")
        else:
            core_properties.comments = ''
        # 保存到桌面文件夹里
        file_name: str = f'D{day_idx + 1}-定点-{factor}'
        safe_file_name: str = re.sub(r'[?*/\<>:"|]', ',', file_name)
        file_output_path: Path = self.output_path/safe_file_name
        doc_copy.save(f'{file_output_path}.docx')

    def write_personnel_deleterious_substance(
            self,
            doc_copy: Any,
            info_dict: Dict[str, Any],
            current_factor_df: DataFrame[Any],
            factor: str,
            day_idx: int,
            schedule: Any
    ):
        '''将个体有害物质信息写入模板'''
        # 计算需要的记录表页数
        table_pages: int = (
            math
            .ceil(
                (len(current_factor_df) - info_dict['first_page_rows'])
                / info_dict['late_page_rows'] + 1
            )
        )

        # 按照页数来增减表格数量
        if table_pages == 1:
            rm_table = doc_copy.tables[2]
            t = rm_table._element
            t.getparent().remove(t)
            rm_page_break = doc_copy.paragraphs[-2]
            pg = rm_page_break._element
            pg.getparent().remove(pg)
            rm_page_break2 = doc_copy.paragraphs[-2]
            pg2 = rm_page_break2._element
            pg2.getparent().remove(pg2)
        elif table_pages == 2:
            pass
        else:
            for _ in range(table_pages - 2):
                cp_table = doc_copy.tables[2]
                new_table = deepcopy(cp_table)
                new_paragraph = doc_copy.add_page_break()
                new_paragraph._p.addnext(new_table._element)
                doc_copy.add_paragraph()
        # 确定不同的表格要填入的信息范围
        tables = doc_copy.tables
        for table_page in range(table_pages):
            if table_page == 0:
                index_first: int = 0
                index_last: int = info_dict['first_page_rows'] - 1
            else:
                index_first: int = (
                    info_dict['late_page_rows']
                    * (table_page - 1)
                    + info_dict['first_page_rows']
                )
                index_last: int = (
                    info_dict['first_page_rows']
                    + table_page
                    * info_dict['late_page_rows']
                    - 1
                )
            current_df: DataFrame[Any] = (
                current_factor_df # type: ignore
                .query(f'index >= {index_first} and index <= {index_last}')
                .reset_index(drop=True)
            )
            # 向指定表格填写数据
            current_table = tables[table_page + 1]
            for r_i in range(current_df.shape[0]):
                row_info = {
                    '采样点编号': current_df.loc[r_i, '测点编号'],
                    '采样岗位': f"{current_df.loc[r_i, '单元']}\n{current_df.loc[r_i, '工种/岗位']}",
                    '样品编号': current_df.loc[r_i, '样品编号'],
                    '代表时长': current_df.loc[r_i, '日接触时长/h'],
                }
                # 采样点编号单元格
                cell1 = current_table.cell(r_i * 3 + 2, 0)
                cell1.text = str(row_info['采样点编号'])
                cell1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                cell1.paragraphs[0].runs[0].font.size = Pt(8)
                # 采样岗位单元格
                cell2 = current_table.cell(r_i * 3 + 2, 1)
                cell2.text = row_info['采样岗位']
                cell2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                cell2.paragraphs[0].runs[0].font.size = Pt(7.5)
                #[x] 样品编号加上项目编号前缀
                # 样品编号单元格
                cell4 = current_table.cell(r_i * 3 + 2, 2)
                cell4.text = f"{self.project_number}{row_info['样品编号']:0>4d}"
                cell4.paragraphs[0].runs[0].font.size = Pt(8)
                # 代表时长
                cell5 = current_table.cell(r_i * 3 + 2, 4)
                cell5.text = str(row_info['代表时长'])
                cell5.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                cell5.paragraphs[0].runs[0].font.size = Pt(9)
        #[x] 写入项目基本信息
        info_table = tables[0]
        # 项目编号
        code_cell = info_table.cell(0, 1)
        code_cell.text = self.project_number
        # 单位
        comp_cell = info_table.cell(0, 4)
        comp_cell.text = self.company_name
        # 检测因素
        item_cell = info_table.cell(3, 1)
        item_cell.text = str(factor)
        # 采样日期
        date_cell = info_table.cell(3, 6)
        if self.schedule_col == '采样/送样日期':
            date_cell.text = schedule.strftime("%Y年%m月%d日")
        for cell in [code_cell, comp_cell, item_cell, date_cell]:
            p = cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # type: ignore
            if len(cell.text) >= 14:
                p.runs[0].font.size = Pt(8)
            else:
                pass
                # p.runs[0].font.size = Pt(9)
        # 页脚信息
        core_properties = doc_copy.core_properties
        core_properties.keywords = factor
        if self.schedule_col == "采样/送样日期":
            core_properties.comments  = schedule.strftime(r"%Y/%m/%d")
        else:
            core_properties.comments = ''
        # 保存到桌面文件夹里
        file_name: str = f'D{day_idx + 1}-个体-{factor}'
        safe_file_name: str = re.sub(r'[?*/\<>:"|]', ',', file_name)
        file_output_path: Path = self.output_path/safe_file_name
        doc_copy.save(f'{file_output_path}.docx')

    def write_personnel_noise(
            self,
            current_factor_df: DataFrame[Any],
            doc: Any,
            current_factor_info: Dict[str, Any],
            day_idx: int = 0,
            schedule: Any = 0
        ) -> None:
        '''将个体噪声信息写入模板'''
        # 读取个体噪声模板
        document = deepcopy(doc)
        # 判断需要的记录表的页数
        table_pages: int = (
            math.ceil(
                (len(current_factor_df) - current_factor_info['first_page_rows'])
                / current_factor_info['late_page_rows']
            )
            + 1
        )
        # 根据不同页数，增减表格
        if table_pages == 1:
            # 删除第二页的表格
            rm_table = document.tables[2]
            t = rm_table._element
            t.getparent().remove(t)
            # 删除最后一个段落
            paragraphs = document.paragraphs
            rm_paragraphs1 = paragraphs[-1]
            rm_p1 = rm_paragraphs1._element
            rm_p1.getparent().remove(rm_p1)
            # 删除倒数第二个段落，即模板的第一页的换页符
            rm_paragraphs2 = paragraphs[-2]
            rm_p2 = rm_paragraphs2._element
            rm_p2.getparent().remove(rm_p2)
        elif table_pages == 2:
            pass # 跳过
        else:
            # 循环增加表格
            for _ in range(table_pages - 2):
                # 复制第二页的表格
                cp_table = document.tables[2]
                new_table = deepcopy(cp_table)
                # 在模板末尾增加段落
                new_paragraph = document.add_page_break()
                # 增加复制的表格
                new_paragraph._p.addnext(new_table._element)
                # 再增加一个段落
                document.add_paragraph()
        # 写入信息
        # 处理后的模板的所有表格
        tables = document.tables
        # 分析不同表格的写入信息
        for table_page in range(table_pages):
            # 获得当前表格的相应信息的索引
            if table_page == 0:
                index_first: int = 0
                index_last: int = current_factor_info['first_page_rows'] - 1
            else:
                index_first: int = (
                    current_factor_info['late_page_rows']
                    * (table_page - 1)
                    + current_factor_info['first_page_rows']
                )
                index_last: int = (
                    current_factor_info['first_page_rows']
                    + table_page
                    * current_factor_info['late_page_rows']
                    - 1
                )
            # 筛选出当前表格的信息
            if index_first == index_last:
                current_df: DataFrame[Any] = (
                    current_factor_df # type: ignore
                    .query(f'index == {index_first}')
                    .reset_index(drop=True)
                )
            else:
                current_table = tables[table_page + 1]
                current_df: DataFrame[Any] = (
                    current_factor_df # type: ignore
                    .query(f'index >= {index_first} and index <= {index_last}')
                    .reset_index(drop=True)
                )
            current_table = tables[table_page + 1]
            # 按行循环选取单元格
            for r_i in range(current_df.shape[0]):
                current_row_list = [
                    current_df.loc[r_i, '测点编号'],
                    f"{current_df.loc[r_i, '单元']} {current_df.loc[r_i, '工种/岗位']}\n",
                    current_df.loc[r_i, '日接触时长/h'],
                ]
                # 再循环列选取单元格，并写入相应信息
                for i, c_i in enumerate(current_factor_info['available_cols']):
                    current_cell = (
                        current_table.rows[
                            r_i * current_factor_info['item_rows']
                            + current_factor_info['title_rows']
                            ]
                        .cells[c_i]
                    )
                    current_cell.text = str(current_row_list[i])
                    current_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    current_cell.paragraphs[0].runs[0].font.size = Pt(6.5)
                    current_cell.paragraphs[0].paragraph_format.line_spacing = Pt(10)
                    # current_cell.paragraphs[0].runs[0].font.name = '宋体'
        info_table = tables[0]
        code_cell = info_table.cell(0, 1)
        comp_cell = info_table.cell(1, 1)

        code_cell.text = self.project_number
        code_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
        comp_cell.text = self.company_name
        comp_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
        # 页脚信息
        date_cell = info_table.cell(
            current_factor_info['date_row'],
            current_factor_info['date_col']
        )
        core_properties = document.core_properties
        # 是否写入采样日期信息
        is_schedule: bool = self.schedule_col == "采样/送样日期" and self.is_all_factors_split
        if is_schedule:
            date_cell.text = schedule.strftime("%Y年%m月%d日")
            date_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
            core_properties.comments  = schedule.strftime(r"%Y/%m/%d")
        else:
            core_properties.comments = ' '
        # [x] 单元格样式
        # 保存到桌面文件夹里
        if self.is_all_factors_split:
            file_name: str = f'D{day_idx + 1}-个体-噪声记录表'
        else:
            file_name: str = '个体噪声记录表'
        safe_file_name: str = re.sub(r'[?*/\<>:"|]', ',', file_name)
        file_output_path: Path = self.output_path/safe_file_name
        document.save(f'{file_output_path}.docx')

    # [x] 将定点仪器直读检测因素的信息写入模板的方法合并
    def write_direct_reading_factors_docx(
            self,
            direct_reading_factor: str,
            day_idx: int = 0,
            schedule: Any = 0
        ) -> None:
        '''将定点仪器直读检测因素的信息写入模板的方法合并'''
        # [x] 去除重复的检测信息
        # 获得检测因素的信息
        factor_key: str = f'{direct_reading_factor}定点'
        current_factor_info: Dict[str, Any] = self.templates_info_dict[factor_key]
        join_char: str = current_factor_info['join_char']
        # 获得检测因素的点位信息
        query_str: str = (
            '采样方式 == "定点"'
            ' and '
            f'检测参数 == "{direct_reading_factor}"'
            ' and '
            '第几个频次 == 1'
        )
        if self.is_all_factors_split:
            current_factor_df: DataFrame = ( # type: ignore
                self.df # type: ignore
                [self.df[self.schedule_col] == schedule]
                .query(query_str)
                .sort_values('测点编号')
                .reset_index(drop=True)
            )
        else:
            current_factor_df: DataFrame = ( # type: ignore
                self.df # type: ignore
                .query(query_str)
                .sort_values('测点编号')
                .reset_index(drop=True)
            )
        # 读取检测因素模板
        current_factor_template: str = self.templates_info_dict[factor_key]['template_path']
        document = Document(current_factor_template)
        # 判断需要的记录表的页数
        table_pages: int = (
            math.ceil(
                (
                    len(current_factor_df) # type: ignore
                    - current_factor_info['first_page_rows']
                )
                / current_factor_info['late_page_rows']
            )
            + 1
        )
        # 根据不同页数，增减表格
        if table_pages == 1:
            # 删除第二页的表格
            rm_table = document.tables[2]
            t = rm_table._element
            t.getparent().remove(t)
            # 删除最后一个段落
            paragraphs = document.paragraphs
            rm_paragraphs1 = paragraphs[-1]
            rm_p1 = rm_paragraphs1._element
            rm_p1.getparent().remove(rm_p1)
            # 删除倒数第二个段落，即模板的第一页的换页符
            rm_paragraphs2 = paragraphs[-3]
            rm_p2 = rm_paragraphs2._element
            rm_p2.getparent().remove(rm_p2)
        elif table_pages == 2:
            pass # 跳过
        else:
            # 循环增加表格
            for _ in range(table_pages - 2):
                # 复制第二页的表格
                cp_table = document.tables[2]
                new_table = deepcopy(cp_table)
                # 在模板末尾增加段落
                new_paragraph = document.add_page_break()
                # 增加复制的表格
                new_paragraph._p.addnext(new_table._element)
                # 再增加一个段落
                document.add_paragraph()
        # [x] 写入信息
        # # 处理后的模板的所有表格
        tables = document.tables
        # 分析不同表格的写入信息
        for table_page in range(table_pages):
            # 获得当前表格的相应信息的索引
            if table_page == 0:
                index_first: int = 0
                index_last: int = current_factor_info['first_page_rows'] - 1
            else:
                index_first: int = (
                    current_factor_info['late_page_rows']
                    * (table_page - 1)
                    + current_factor_info['first_page_rows']
                )
                index_last: int = (
                    current_factor_info['first_page_rows']
                    + table_page
                    * current_factor_info['late_page_rows']
                    - 1
                )
            # 筛选出当前表格的信息
            if index_first == index_last:
                current_df: DataFrame = ( # type: ignore
                    current_factor_df # type: ignore
                    .query(f'index == {index_first}')
                    .reset_index(drop=True)
                )
            else:
                current_table = tables[table_page + 1]
                current_df: DataFrame = ( # type: ignore
                    current_factor_df # type: ignore
                    .query(f'index >= {index_first} and index <= {index_last}')
                    .reset_index(drop=True)
                )
            current_table = tables[table_page + 1]
            # 按行循环选取单元格
            for r_i in range(current_df.shape[0]): # type: ignore
                current_row_list = [ # type: ignore
                    current_df.loc[r_i, '测点编号'], # type: ignore
                    f"{current_df.loc[r_i, '单元']}{join_char}{current_df.loc[r_i, '检测地点']}", # type: ignore
                    current_df.loc[r_i, '日接触时长/h'], # type: ignore
                ]
                # 再循环列选取单元格，并写入相应信息
                for i, c_i in enumerate(current_factor_info['available_cols']):
                    current_cell = (
                        current_table.rows[
                            r_i * current_factor_info['item_rows']
                            + current_factor_info['title_rows']
                            ]
                        .cells[c_i]
                    )
                    current_cell.text = str(current_row_list[i]) # type: ignore
                    current_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    current_cell.paragraphs[0].runs[0].font.size = Pt(8)
                    current_cell.paragraphs[0].paragraph_format.line_spacing = Pt(10)
                    # current_cell.paragraphs[0].runs[0].font.name = '宋体'
        # [x] 样式调整
        # 写入基本信息
        info_table = tables[0]
        code_cell = (
            info_table
            .rows[current_factor_info['project_num_row']]
            .cells[current_factor_info['project_num_col']]
        )
        comp_cell = (
            info_table
            .rows[current_factor_info['company_name_row']]
            .cells[current_factor_info['company_name_col']]
        )
        code_cell.text = self.project_number
        code_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
        comp_cell.text = self.company_name
        comp_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
        # [x] 日期信息
        date_cell = info_table.cell(
            current_factor_info['date_row'],
            current_factor_info['date_col']
        )
        is_schedule: bool = self.schedule_col == "采样/送样日期" and self.is_all_factors_split
        # 页脚信息
        core_properties = document.core_properties
        if is_schedule:
            date_cell.text = schedule.strftime("%Y年%m月%d日")
            date_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
            core_properties.comments  = schedule.strftime(r"%Y/%m/%d")
        else:
            core_properties.comments = ''
        # 保存到桌面文件夹里
        if self.is_all_factors_split:
            file_name: str = f'D{day_idx + 1}-直读-{direct_reading_factor}记录表'
        else:
            file_name: str = f'直读-{factor_key}记录表'
        safe_file_name: str = re.sub(r'[?*/\<>:"|]', ',', file_name)
        file_output_path: Path = self.output_path/safe_file_name
        document.save(f'{file_output_path}.docx')

    def write_traveler_docx(
            self,
            current_traveler_df: DataFrame[Any],
            doc: Any,
            day_i: int,
            schedule: Any
        ) -> None:
        '''将流转单信息写入模板'''
        project_num_cell = doc.tables[0].rows[0].cells[1]
        project_num_cell.text = self.project_number
        # [x] 样式
        # 判断需要的流转单的页数
        table_pages: int = math.ceil(len(current_traveler_df) / 8)
        for _ in range(table_pages - 1):
            cp_table = doc.tables[0]
            new_table = deepcopy(cp_table)
            cp_paragraph = doc.paragraphs[0]
            last_paragraph = doc.add_page_break()
            last_paragraph._p.addnext(new_table._element)
            doc.add_paragraph(cp_paragraph.text)

        tables = doc.tables
        # 写入各个表格里
        for table_page in range(table_pages):
            first_index: int = 8 * table_page
            last_index: int = 8 * table_page + 7
            # .reset_index(drop=True)
            current_df: DataFrame[Any] = (
                current_traveler_df # type: ignore
                .iloc[first_index : last_index + 1]
                .reset_index()
            )
            current_table = tables[table_page]
            # 获得每行的信息
            for r_i in range(len(current_df)):
                num_range_str: str = ( # type: ignore
                    ','.join(current_df.loc[r_i, "全部样品编号"])  # type: ignore
                )
                current_row_list = [
                    f'{self.project_number}{num_range_str}',
                    current_df.loc[r_i, "检测参数"],  # type: ignore
                    current_df.loc[r_i, "保存时间"],  # type: ignore
                    current_df.loc[r_i, "全部样品数量"],  # type: ignore
                ]
                for c_i in list(range(4)):
                    match_cols_list: List[int] = [0, 1, 3, 4]
                    current_cell = (
                        current_table
                        .rows[r_i + 2]
                        .cells[match_cols_list[c_i]]
                    )
                    current_cell.text = str(current_row_list[c_i])
                    current_cell.paragraphs[0].runs[0].font.size = Pt(7.5)
                    current_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    # [x] 单元格样式
                    if '\\n' in current_cell.text:
                        new_text: str = current_cell.text.replace('\\n', '\n')
                        current_cell.text = new_text
                        current_cell.paragraphs[0].runs[0].font.size = Pt(7.5)
                        current_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
        # 页脚信息
        core_properties = doc.core_properties
        if self.schedule_col == "采样/送样日期":
            core_properties.comments  = schedule.strftime(r"    %Y年  %m月  %d日")
        else:
            core_properties.comments = '        年    月    日'
        # 保存到桌面文件夹里
        file_name: str = f'D{day_i + 1}-样品流转单'
        safe_file_name: str = re.sub(r'[?*/\<>:"|]', ',', file_name)
        file_output_path: Path = self.output_path/safe_file_name
        doc.save(f'{file_output_path}.docx')
