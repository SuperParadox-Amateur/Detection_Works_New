'''将信息写入模板文件里'''
from nptyping import DataFrame
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

def write_point_deleterious_substance(
        doc_info
        doc: Any, day_i: int, schedule: Any
    ):
    '''将定点有害物质信息写入模板'''
    # merger = PdfWriter()
    # for day_i, schedule in enumerate(self.schedule_list):
    today_df = (
            self
            .point_df
            [self.point_df[self.schedule_col] == schedule]
            # .query(f'{self.schedule_col} == @schedule')
            .sort_values(by=['测点编号'])
            .reset_index(drop=True)
    )
    # 采样日期
    # schedule_str = datetime.fromtimestamp(schedule).strftime("%Y-%m-%d")
    # schedule_dt = datetime.strptime(schedule, '%Y-%m-%d') # type: ignore
    schedule_dt = datetime.strptime(schedule, r'%Y/%m/%d') # type: ignore
    factors: List[str] = today_df['检测参数'].drop_duplicates().tolist()
    sorted_factors: List[str] = sorted(factors, key=lambda x: x.encode('gbk'))
    # 获得当前检测因素的dataframe
    for factor in sorted_factors:
        # 导入定点模板
        doc_copy = deepcopy(doc)
        # 获得当前检测因素的dataframe
        current_factor_df = (
            today_df[today_df['检测参数'] == factor]
            .sort_values(by='测点编号')
            .reset_index(drop=True)
        )
        # 计算需要的记录表页数
        table_pages: int = (
            math
            .ceil(
                (len(current_factor_df) - 3)
                / 4 + 1
            )
        )
        # 按照页数来增减表格数量
        if table_pages == 1:
            rm_table = doc_copy.tables[2]
            t = rm_table._element
            t.getparent().remove(t)
            rm_page_break = doc_copy.paragraphs[-2]
            pg = rm_page_break._element
            pg.getparent().remove(pg)
            rm_page_break2 = doc_copy.paragraphs[-2]
            pg2 = rm_page_break2._element
            pg2.getparent().remove(pg2)
        elif table_pages == 2:
            pass
        else:
            for _ in range(table_pages - 2):
                cp_table = doc_copy.tables[2]
                new_table = deepcopy(cp_table)
                new_paragraph = doc_copy.add_page_break()
                new_paragraph._p.addnext(new_table._element)
                doc_copy.add_paragraph()
        # 确定不同的表格要填入的信息范围
        tables = doc_copy.tables
        for table_page in range(table_pages):
            if table_page == 0:
                index_first: int = 0
                index_last: int = 2
            else:
                index_first: int = 4 * table_page - 1
                index_last: int = 4 * table_page + 2
            current_df = (
                current_factor_df
                .query(f'index >= {index_first} and index <= {index_last}')
                .reset_index(drop=True)
            )
            # 向指定表格填写数据
            current_table = tables[table_page + 1]
            for r_i in range(current_df.shape[0]):
                row_info = {
                    '采样点编号': current_df.loc[r_i, '测点编号'],
                    '采样岗位': f"{current_df.loc[r_i, '单元']}\n{current_df.loc[r_i, '检测地点']}",
                    '空白编号1': current_df.loc[r_i, '空白编号1'],
                    '空白编号2': current_df.loc[r_i, '空白编号2'],
                    '样品编号': current_df.loc[r_i, '样品编号'],
                    '代表时长': current_df.loc[r_i, '代表时长'],
                    '是否合并代表时长': current_df.loc[r_i, '是否合并代表时长'],
                }
                # 采样点编号单元格
                cell1 = current_table.cell(r_i * 6 + 2, 0)
                cell1.text = str(row_info['采样点编号'])
                cell1.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                cell1.paragraphs[0].runs[0].font.size = Pt(8)
                # 采样岗位单元格
                cell2 = current_table.cell(r_i * 6 + 2, 1)
                cell2.text = row_info['采样岗位']
                cell2.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                cell2.paragraphs[0].runs[0].font.size = Pt(7.5)
                #[x] 样品编号加上项目编号前缀
                # 空白编号单元格，只写入第一行
                if table_page == 0 and r_i == 0 and row_info['空白编号1'] != '-':
                    cell3_1 = current_table.cell(r_i * 6 + 2, 2)
                    cell3_1.text = f"{self.project_number}{row_info['空白编号1']}"
                    cell3_2 = current_table.cell(r_i * 6 + 3, 2)
                    cell3_2.text = f"{self.project_number}{row_info['空白编号2']}"
                    cell3_1.paragraphs[0].runs[0].font.size = Pt(8)
                    cell3_2.paragraphs[0].runs[0].font.size = Pt(8)
                else:
                    pass
                # 样品编号单元格
                for n_i, num in enumerate(row_info['样品编号']):
                    cell4 = current_table.cell(r_i * 6 + n_i + 4, 2)
                    cell4.text = f"{self.project_number}{num:0>4d}"
                    cell4.paragraphs[0].runs[0].font.size = Pt(8)
                # 代表时长
                for n_i, duration in enumerate(row_info['代表时长']):
                    cell5 = current_table.cell(r_i * 6 + n_i + 4, 9)
                    cell5.text = str(duration)
                    cell5.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # type: ignore
                    cell5.paragraphs[0].runs[0].font.size = Pt(9)
                # 是否合并代表时长
                if row_info['是否合并代表时长'] == True:
                    merge_len: int = len(row_info['样品编号'])
                    merge_cell1 = current_table.cell(r_i * 6 + 4, 9)
                    merge_cell2 = current_table.cell(r_i * 6 + merge_len + 3, 9)
                    merge_cell1.merge(merge_cell2)
                #[x] 样式调整
        #[x] 写入项目基本信息
        info_table = tables[0]
        # 项目编号
        code_cell = info_table.cell(0, 1)
        code_cell.text = self.project_number
        # 单位
        comp_cell = info_table.cell(0, 4)
        comp_cell.text = self.company_name
        # 检测因素
        item_cell = info_table.cell(3, 1)
        item_cell.text = str(factor)
        # 采样日期
        date_cell = info_table.cell(3, 6)
        if self.schedule_col == '采样/送样日期':
            date_cell.text = schedule_dt.strftime("%Y年%m月%d日")
        for cell in [code_cell, comp_cell, item_cell, date_cell]:
            p = cell.paragraphs[0]
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # type: ignore
            if len(cell.text) >= 14:
                p.runs[0].font.size = Pt(8)
            else:
                pass
                # p.runs[0].font.size = Pt(9)
        # 页脚信息
        core_properties = doc_copy.core_properties
        core_properties.keywords = factor
        if self.schedule_col == "采样/送样日期":
            core_properties.comments  = schedule_dt.strftime(r"%Y/%m/%d")
        # 保存到桌面文件夹里
        file_name = f'D{day_i + 1}-定点-{factor}'
        safe_file_name: str = re.sub(r'[?*/\<>:"|]', ',', file_name)
        file_output_path: str = os.path.join(
            self.output_path,
            safe_file_name
        )
        doc_copy.save(f'{file_output_path}.docx')